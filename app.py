import streamlit as st
import os
import json
import pandas as pd
from datetime import datetime, date
import calendar
import io
import mimetypes
import re
import tempfile
import time
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional

from google import genai
from google.genai import types
from pydantic import BaseModel, ConfigDict, Field, ValidationError
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENTATION
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from fpdf import FPDF, XPos, YPos
from github import Github, GithubException, UnknownObjectException  # <-- Added for GitHub backups

# --- 1. CONFIGURATION & SETUP ---
st.set_page_config(page_title="AINP Grant Manager", page_icon="🌾", layout="wide")

# Directory Setup
DIRS = ['data', 'documents', 'fonts', 'logos']
for d in DIRS:
    if not os.path.exists(d):
        os.makedirs(d)

def get_secret_value(*names, default=""):
    """Read Streamlit secrets safely, with environment variables as a local fallback."""
    for name in names:
        value = None
        try:
            value = st.secrets.get(name)
        except Exception:
            value = None
        if not value:
            value = os.environ.get(name)
        if value:
            return str(value).strip()
    return default

DEFAULT_GITHUB_REPO = "vkcvaibhav/AINP-Grant-Manager"
GEMINI_HEAVY_MODEL = get_secret_value("GEMINI_HEAVY_MODEL", "GEMINI_PRO_MODEL", "GEMINI_MAIN_MODEL", default="gemini-3.1-pro-preview")
GEMINI_CHAT_MODEL = get_secret_value("GEMINI_CHAT_MODEL", default="gemini-3.5-flash")
GEMINI_FAST_MODEL = get_secret_value("GEMINI_FAST_MODEL", default="gemini-3.1-flash-lite")
GEMINI_EMBEDDING_MODEL = get_secret_value("GEMINI_EMBEDDING_MODEL", default="models/gemini-embedding-2")
AI_KNOWLEDGE_FILE = 'data/ai_knowledge.json'
LEARNING_MEMORY_FILE = 'data/learning_memory.json'

# AI Setup
api_key = get_secret_value("GEMINI_API_KEY", "GOOGLE_API_KEY")
genai_client = genai.Client(api_key=api_key) if api_key else None

# Load Logos if exist
NAU_LOGO = 'logos/nau_logo.png' if os.path.exists('logos/nau_logo.png') else None
ICAR_LOGO = 'logos/icar_logo.png' if os.path.exists('logos/icar_logo.png') else None
GUJARATI_FONT = 'fonts/NotoSansGujarati-Regular.ttf'

# Define standard Heads
BUDGET_HEADS = [
    "Pay and Allowances",
    "Travelling Allowances (TA)",
    "Other Recurring Contingencies (ORC)",
    "Non-Recurring Contingencies (Equipments/Works)",
    "TSP"
]

# --- 2. HELPER FUNCTIONS ---

class BudgetHeadExtraction(BaseModel):
    head_name: str = ""
    icar_share: float = 0.0
    state_share: float = 0.0
    total: float = 0.0


class BudgetExtraction(BaseModel):
    is_revision: bool = False
    date: Optional[str] = None
    heads: List[BudgetHeadExtraction] = Field(default_factory=list)


class InstallmentHeadExtraction(BaseModel):
    head_name: str = ""
    amount: float = 0.0


class InstallmentExtraction(BaseModel):
    date: Optional[str] = None
    installment_number: str = ""
    purpose: str = ""
    pfms_transaction_id: str = ""
    heads: List[InstallmentHeadExtraction] = Field(default_factory=list)


class AucBalanceExtraction(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    establishment_charges: float = Field(0.0, alias="Establishment Charges")
    ta: float = Field(0.0, alias="TA")
    contingencies: float = Field(0.0, alias="Contingencies")
    tsp: float = Field(0.0, alias="TSP")
    equipments: float = Field(0.0, alias="Equipments")
    works: float = Field(0.0, alias="Works")


def coerce_amount(value, default=0.0):
    try:
        default_value = float(default)
    except (TypeError, ValueError):
        default_value = 0.0
    if value is None:
        return default_value
    if isinstance(value, bool):
        return default_value
    if isinstance(value, (int, float)):
        try:
            if pd.isna(value):
                return default_value
        except Exception:
            pass
        return float(value)
    text = str(value).strip().replace(",", "")
    match = re.search(r"-?\d+(?:\.\d+)?", text)
    if not match:
        return default_value
    try:
        return float(match.group(0))
    except ValueError:
        return default_value


def safe_date_string(value, fallback=None):
    fallback = fallback or date.today().strftime("%Y-%m-%d")
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")
    if value:
        try:
            return pd.to_datetime(value, errors="raise").strftime("%Y-%m-%d")
        except Exception:
            return fallback
    return fallback


def parse_date(value):
    try:
        return datetime.strptime(safe_date_string(value), "%Y-%m-%d")
    except Exception:
        return None


def get_fy_start_date(fy):
    try:
        return f"{int(str(fy).split('-')[0])}-04-01"
    except Exception:
        return date.today().strftime("%Y-%m-%d")


def normalize_budget_map(raw_map):
    normalized = {}
    if not isinstance(raw_map, dict) or not raw_map:
        return normalized
    for head in BUDGET_HEADS:
        values = raw_map.get(head, {})
        if not isinstance(values, dict):
            values = {}
        total = coerce_amount(values.get('total'))
        icar = coerce_amount(values.get('icar'))
        state = coerce_amount(values.get('state'))
        normalized[head] = {'icar': icar, 'state': state, 'total': total}
    return normalized


def normalize_fy_data(raw_data, fy):
    data = get_default_data(fy)
    if isinstance(raw_data, dict):
        data.update(raw_data)
    data['financial_year'] = fy
    data['allocation'] = normalize_budget_map(data.get('allocation', {}))
    data['revised_allocation'] = normalize_budget_map(data.get('revised_allocation', {}))
    data['quarterly_allocations'] = data.get('quarterly_allocations') if isinstance(data.get('quarterly_allocations'), dict) else {}
    data['quarterly_allocations'] = {
        q: normalize_budget_map(data['quarterly_allocations'].get(q, {}))
        for q in ["Q1", "Q2", "Q3", "Q4"]
    }

    normalized_installments = []
    for inst in data.get('installments', []):
        if not isinstance(inst, dict):
            continue
        item = dict(inst)
        item['date'] = safe_date_string(item.get('date'), get_fy_start_date(fy))
        item['pfms_id'] = str(item.get('pfms_id') or item.get('pfms_transaction_id') or f"UNKNOWN_{len(normalized_installments)+1}").strip()
        item['installment_num'] = str(item.get('installment_num') or item.get('installment_number') or item.get('type') or "").strip()
        item['type'] = str(item.get('type') or item['installment_num'] or item['pfms_id']).strip()
        item['purpose'] = str(item.get('purpose') or "").strip()
        item['quarter'] = item.get('quarter') if item.get('quarter') in ["Q1", "Q2", "Q3", "Q4"] else quarter_from_date(item['date'])
        heads = item.get('heads') if isinstance(item.get('heads'), dict) else {}
        item['heads'] = {head: coerce_amount(heads.get(head)) for head in BUDGET_HEADS}
        item['amount'] = coerce_amount(item.get('amount'), sum(item['heads'].values()))
        item['available'] = bool(item.get('available', False))
        item['comptroller_order_uploaded'] = bool(item.get('comptroller_order_uploaded', False))
        normalized_installments.append(item)
    data['installments'] = normalized_installments

    normalized_expenses = []
    for exp in data.get('expenditure', []):
        if not isinstance(exp, dict):
            continue
        item = dict(exp)
        item['date'] = safe_date_string(item.get('date'), get_fy_start_date(fy))
        item['head'] = str(item.get('head') or BUDGET_HEADS[0])
        item['detail'] = str(item.get('detail') or "")
        item['amount'] = coerce_amount(item.get('amount'))
        item.pop('_orig_idx', None)
        normalized_expenses.append(item)
    data['expenditure'] = normalized_expenses

    ob = data.get('opening_balances') if isinstance(data.get('opening_balances'), dict) else {}
    data['opening_balances'] = {
        "Establishment Charges": coerce_amount(ob.get("Establishment Charges")),
        "TA": coerce_amount(ob.get("TA")),
        "Contingencies": coerce_amount(ob.get("Contingencies")),
        "TSP": coerce_amount(ob.get("TSP")),
        "Equipments": coerce_amount(ob.get("Equipments")),
        "Works": coerce_amount(ob.get("Works")),
    }
    data['latest_quarter'] = str(data.get('latest_quarter') or "Full Year")
    data['latest_date'] = str(data.get('latest_date') or "N/A")
    return data


def quarter_from_date(value):
    parsed = parse_date(value)
    month = parsed.month if parsed else date.today().month
    if month in [4, 5, 6]:
        return "Q1"
    if month in [7, 8, 9]:
        return "Q2"
    if month in [10, 11, 12]:
        return "Q3"
    return "Q4"


def sanitize_filename(filename, default="file"):
    name = os.path.basename(str(filename or default)).strip()
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", name)
    name = re.sub(r"\s+", " ", name).strip(" .")
    return name or default


def safe_documents_path(filename):
    safe_name = sanitize_filename(filename)
    docs_dir = os.path.abspath("documents")
    path = os.path.abspath(os.path.join(docs_dir, safe_name))
    if not path.startswith(docs_dir + os.sep):
        raise ValueError("Unsafe document path")
    return path


def save_uploaded_document(uploaded_file, filename):
    path = safe_documents_path(filename)
    with open(path, "wb") as f:
        f.write(uploaded_file.getvalue())
    backup_file_to_github(path)
    return path


def read_json_file(path, default):
    if not os.path.exists(path):
        return default
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, OSError):
        return default


def write_json_file(path, data):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    tmp_path = f"{path}.tmp"
    with open(tmp_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    os.replace(tmp_path, path)
    backup_file_to_github(path)


def get_audit_filename(fy):
    return f"data/audit_log_{fy.replace('-', '_')}.jsonl"


def append_audit_log(fy, action, details=None):
    entry = {
        "timestamp": datetime.now().isoformat(timespec="seconds"),
        "financial_year": fy,
        "action": action,
        "details": details or {},
    }
    path = get_audit_filename(fy)
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "a", encoding="utf-8") as f:
        f.write(json.dumps(entry, ensure_ascii=False) + "\n")
    backup_file_to_github(path)


def read_audit_entries(fy, limit=20):
    path = get_audit_filename(fy)
    entries = []
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                try:
                    entries.append(json.loads(line))
                except json.JSONDecodeError:
                    continue
    return list(reversed(entries[-limit:]))


def build_fy_backup_zip(fy, data):
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        data_name = get_fy_filename(fy)
        zf.writestr(data_name, json.dumps(normalize_fy_data(data, fy), indent=4, ensure_ascii=False))
        audit_name = get_audit_filename(fy)
        if os.path.exists(audit_name):
            zf.write(audit_name, audit_name)
        for extra in [AI_KNOWLEDGE_FILE, LEARNING_MEMORY_FILE]:
            if os.path.exists(extra):
                zf.write(extra, extra)
        docs_dir = Path("documents")
        if docs_dir.exists():
            fy_safe = fy.replace("-", "_")
            for path in docs_dir.glob("*"):
                if path.is_file() and (fy in path.name or fy_safe in path.name or path.name.startswith("AUC_Archive_")):
                    zf.write(path, str(path).replace("\\", "/"))
    buffer.seek(0)
    return buffer


def preview_restore_json(uploaded_file, fy):
    try:
        raw = uploaded_file.getvalue().decode("utf-8")
        parsed = json.loads(raw)
        restored = normalize_fy_data(parsed, fy)
        return restored, None
    except Exception as e:
        return None, str(e)


def pdf_to_bytes(pdf):
    raw = pdf.output()
    if isinstance(raw, bytearray):
        return bytes(raw)
    if isinstance(raw, bytes):
        return raw
    return raw.encode("latin-1")


def load_ai_knowledge():
    return read_json_file(AI_KNOWLEDGE_FILE, {"store_name": "", "display_name": "", "documents": []})


def save_ai_knowledge(metadata):
    write_json_file(AI_KNOWLEDGE_FILE, metadata)


def ensure_file_search_store():
    if genai_client is None:
        raise RuntimeError("GEMINI_API_KEY is required for File Search.")
    metadata = load_ai_knowledge()
    store_name = metadata.get("store_name")
    if store_name:
        try:
            genai_client.file_search_stores.get(name=store_name)
            return metadata
        except Exception:
            metadata["store_name"] = ""
    store = genai_client.file_search_stores.create(
        config=types.CreateFileSearchStoreConfig(
            display_name="AINP Grant Manager Knowledge Store",
            embedding_model=GEMINI_EMBEDDING_MODEL,
        )
    )
    metadata["store_name"] = store.name
    metadata["display_name"] = getattr(store, "display_name", None) or "AINP Grant Manager Knowledge Store"
    save_ai_knowledge(metadata)
    return metadata


def upload_file_to_knowledge_store(uploaded_file):
    metadata = ensure_file_search_store()
    original_name = sanitize_filename(uploaded_file.name, "knowledge_file")
    mime_type = uploaded_file.type or mimetypes.guess_type(original_name)[0] or "application/octet-stream"
    with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{original_name}") as tmp:
        tmp.write(uploaded_file.getvalue())
        tmp_path = tmp.name
    try:
        operation = genai_client.file_search_stores.upload_to_file_search_store(
            file_search_store_name=metadata["store_name"],
            file=tmp_path,
            config=types.UploadToFileSearchStoreConfig(
                display_name=original_name,
                mime_type=mime_type,
            ),
        )
        started = time.time()
        while not operation.done and time.time() - started < 180:
            time.sleep(3)
            operation = genai_client.operations.get(operation)
        doc_info = operation.model_dump(mode="json") if hasattr(operation, "model_dump") else {"name": str(operation)}
        metadata.setdefault("documents", []).append({
            "file_name": original_name,
            "mime_type": mime_type,
            "uploaded_at": datetime.now().isoformat(timespec="seconds"),
            "operation": doc_info,
        })
        save_ai_knowledge(metadata)
        return operation
    finally:
        try:
            os.remove(tmp_path)
        except OSError:
            pass


def load_learning_memories():
    raw = read_json_file(LEARNING_MEMORY_FILE, [])
    return raw if isinstance(raw, list) else []


def save_learning_memories(memories):
    write_json_file(LEARNING_MEMORY_FILE, memories)


def matching_learning_memories(prompt):
    prompt_l = str(prompt or "").lower()
    matches = []
    for memory in load_learning_memories():
        if not memory.get("enabled", True):
            continue
        haystack = " ".join([
            str(memory.get("title", "")),
            str(memory.get("keywords", "")),
            str(memory.get("memory_text", "")),
        ]).lower()
        keywords = [k.strip().lower() for k in str(memory.get("keywords", "")).split(",") if k.strip()]
        if any(k in prompt_l for k in keywords) or any(word and word in prompt_l for word in str(memory.get("title", "")).lower().split()):
            matches.append(memory)
        elif not keywords and haystack and any(word in haystack for word in prompt_l.split()[:8]):
            matches.append(memory)
    return matches[:5]


def extract_file_search_citations(response):
    citations = []
    for candidate in response.candidates or []:
        metadata = getattr(candidate, "grounding_metadata", None)
        if not metadata or not metadata.grounding_chunks:
            continue
        for chunk in metadata.grounding_chunks:
            chunk_data = chunk.model_dump(mode="json") if hasattr(chunk, "model_dump") else {}
            citations.append(chunk_data)
    return citations

# --- A. Data Persistence & GitHub Backup ---
def get_github_repo():
    github_token = get_secret_value("GITHUB_TOKEN")
    repo_name = get_secret_value("GITHUB_REPO", "REPO_NAME", default=DEFAULT_GITHUB_REPO)
    if not github_token:
        return None
    g = Github(github_token)
    return g.get_repo(repo_name)


def backup_to_github(filepath, content):
    """Pushes saved data back to GitHub so it isn't lost when Streamlit sleeps."""
    if not get_secret_value("GITHUB_TOKEN"):
        # If no token is found, just skip the backup (useful for local testing)
        return

    try:
        repo = get_github_repo()
        if repo is None:
            return
        
        # Check if the file already exists in GitHub
        try:
            contents = repo.get_contents(filepath)
        except UnknownObjectException:
            # If it does not exist yet, CREATE it
            repo.create_file(filepath, f"Auto-create {filepath}", content)
        else:
            if isinstance(contents, list):
                st.warning(f"GitHub backup skipped because {filepath} resolved to a directory.")
                return
            # If it exists, UPDATE it
            repo.update_file(contents.path, f"Auto-backup {filepath}", content, contents.sha)
            
    except GithubException as e:
        st.warning(f"Failed to backup to GitHub. Error: {e}")
    except Exception as e:
        st.warning(f"Failed to backup to GitHub. Error: {e}")


def backup_file_to_github(local_path, github_path=None):
    if not get_secret_value("GITHUB_TOKEN") or not os.path.exists(local_path):
        return
    if github_path is None:
        try:
            github_path = os.path.relpath(local_path, os.getcwd())
        except ValueError:
            github_path = local_path
    github_path = github_path.replace("\\", "/")
    try:
        mode = "rb" if not local_path.endswith((".json", ".jsonl", ".txt", ".py", ".md")) else "r"
        if mode == "rb":
            with open(local_path, "rb") as f:
                content = f.read()
        else:
            with open(local_path, "r", encoding="utf-8") as f:
                content = f.read()
        backup_to_github(github_path, content)
    except Exception as e:
        st.warning(f"Failed to backup file to GitHub. Error: {e}")


def pull_file_from_github(github_path, local_path):
    if os.path.exists(local_path) or not get_secret_value("GITHUB_TOKEN"):
        return os.path.exists(local_path)
    try:
        repo = get_github_repo()
        if repo is None:
            return False
        contents = repo.get_contents(github_path.replace("\\", "/"))
        if isinstance(contents, list):
            return False
        os.makedirs(os.path.dirname(local_path), exist_ok=True)
        with open(local_path, "wb") as f:
            f.write(contents.decoded_content)
        return True
    except Exception:
        return False


def load_document_bytes(file_path):
    if os.path.exists(file_path):
        with open(file_path, "rb") as f:
            return f.read()
    try:
        github_path = os.path.relpath(file_path, os.getcwd()).replace("\\", "/")
    except ValueError:
        github_path = file_path.replace("\\", "/")
    if pull_file_from_github(github_path, file_path):
        with open(file_path, "rb") as f:
            return f.read()
    return None

def get_fy_filename(fy):
    return f'data/grant_data_{fy.replace("-", "_")}.json'

def get_default_data(fy):
    return {
        "financial_year": fy,
        "allocation": {},
        "revised_allocation": {},
        "quarterly_allocations": {"Q1": {}, "Q2": {}, "Q3": {}, "Q4": {}},
        "installments": [],
        "expenditure": [],
        "latest_quarter": "Full Year",
        "latest_date": "N/A"
    }

def load_data(fy):
    filename = get_fy_filename(fy)
    if os.path.exists(filename):
        try:
            with open(filename, 'r', encoding='utf-8') as f:
                return normalize_fy_data(json.load(f), fy)
        except (json.JSONDecodeError, OSError) as e:
            st.warning(f"Could not load {filename}; starting with a clean FY file. Error: {e}")
    return normalize_fy_data(get_default_data(fy), fy)

def save_data(data, fy, audit_action=None, audit_details=None):
    filename = get_fy_filename(fy)
    data = normalize_fy_data(data, fy)
    
    # 1. Save locally for immediate app use
    temp_filename = f"{filename}.tmp"
    with open(temp_filename, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=4, ensure_ascii=False)
    os.replace(temp_filename, filename)
        
    # 2. Push to GitHub for permanent backup
    json_str = json.dumps(data, indent=4, ensure_ascii=False)
    backup_to_github(filename, json_str)
    if audit_action:
        append_audit_log(fy, audit_action, audit_details)


# --- B. Document Processing (AI PDF Reader) ---
def extract_structured_with_ai(parts, prompt, schema_model):
    if genai_client is None:
        st.error("Please add GEMINI_API_KEY in Streamlit secrets to use AI extraction.")
        return None
    try:
        response = genai_client.models.generate_content(
            model=GEMINI_HEAVY_MODEL,
            contents=[types.Part.from_text(text=prompt)] + parts,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=schema_model,
            ),
        )
        if response.parsed is not None:
            if isinstance(response.parsed, BaseModel):
                return response.parsed.model_dump(by_alias=True)
            return response.parsed
        return schema_model.model_validate_json(response.text).model_dump(by_alias=True)
    except (ValidationError, json.JSONDecodeError) as e:
        st.error(f"AI returned data in an unexpected format: {e}")
        return None
    except Exception as e:
        st.error(f"Error processing document with AI: {e}")
        return None


def process_upload_with_ai(uploaded_file, prompt_task):
    """Uses Gemini to natively read the PDF and extract structured data."""
    pdf_part = types.Part.from_bytes(data=uploaded_file.getvalue(), mime_type="application/pdf")
    structure = prompt_task.get("structure", {})
    schema_model = BudgetExtraction if "heads" in structure else AucBalanceExtraction
    full_prompt = f"""
    Analyze the attached document content and extract the required information.

    DOCUMENT TYPE/CONTEXT: {prompt_task['context']}

    IMPORTANT:
    - Convert dates to YYYY-MM-DD format.
    - Convert currency values to numbers only, in absolute rupees.
    - If a field cannot be found, set it to null or zero as appropriate.
    """
    return extract_structured_with_ai([pdf_part], full_prompt, schema_model)

# --- C. Output Generation (PDF/WORD) ---
def add_bottom_border(paragraph, size='24'):
    """Adds a bottom border to a paragraph. Size '24' is a thick 3pt line, '8' is a thin 1pt line."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size) 
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)

def generate_comptroller_docx(ref_no, letter_date, body_text, amt_words, pay_amt, total_rec, non_rec_amt, total_amt):
    """Generates the Native Microsoft Word format for the Comptroller Letter matching the exact layout."""
    doc = Document()
    
    # Set document margins tight to match letterhead
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Strip default spacing to ensure tables pack tightly together
    style = doc.styles['Normal']
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
        
    # 1. Header Table for Logos and Center Text
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(3.6)
    table.columns[2].width = Inches(1.4)
    
    # Left Logo (NAU)
    if 'NAU_LOGO' in globals() and NAU_LOGO and os.path.exists(NAU_LOGO):
        cell_left = table.cell(0, 0)
        cell_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        p_left = cell_left.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_left.paragraph_format.space_after = Pt(0)
        p_left.paragraph_format.line_spacing = 0.5
        
        r_left = p_left.add_run()
        r_left.add_picture(NAU_LOGO, width=Inches(1.8))
        
    # Center Text
    p_center = table.cell(0, 1).paragraphs[0]
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_center.paragraph_format.line_spacing = 0.85 # Tight gap below the main heading
    
    r1 = p_center.add_run("કીટકશાસ્ત્ર વિભાગ\n")
    r1.bold = True
    r1.font.size = Pt(22) # Large, prominent heading
    
    r2 = p_center.add_run("ન. મ. કૃષિ મહાવિદ્યાલય\nનવસારી કૃષિ યુનિવર્સિટી\nનવસારી- ૩૯૬ ૪૫૦ (ગુજરાત)")
    r2.bold = True
    r2.font.size = Pt(14)
    
    # Right Logo (ICAR)
    if 'ICAR_LOGO' in globals() and ICAR_LOGO and os.path.exists(ICAR_LOGO):
        p_right = table.cell(0, 2).paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_right = p_right.add_run()
        r_right.add_picture(ICAR_LOGO, width=Inches(1.5))
        
    # Draw First Thick Black Separator Line (Squashed height to remove gap)
    p_thick1 = doc.add_paragraph()
    p_thick1.paragraph_format.space_before = Pt(0)
    p_thick1.paragraph_format.space_after = Pt(0)
    p_thick1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p_thick1.paragraph_format.line_spacing = Pt(1)
    p_thick1.add_run().font.size = Pt(1) 
    add_bottom_border(p_thick1, size='24')
    
    # 2. Sender Info Block
    table2 = doc.add_table(rows=1, cols=2)
    table2.autofit = False
    table2.columns[0].width = Inches(3.4)
    table2.columns[1].width = Inches(3.4)
    
    p1 = table2.cell(0,0).paragraphs[0]
    p1.add_run("ડૉ. જે. જે. પસ્તાગિયા\nપ્રાધ્યાપક અને વડા (ઈ/ચા.)")
    
    p2 = table2.cell(0,1).paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.add_run("મોબાઇલ: +૯૧ ૯૮૭૯૦ ૩૮૫૩૯\nઇમેલ: headentonau@gmail.com")
    
    # Ensure inner paragraphs of table2 have zero spacing
    for cell in table2.rows[0].cells:
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 0.8  # <--- This line squashes the \n gap
    
    # Draw Second Thick Black Separator Line (Squashed height to remove gap)
    p_thick2 = doc.add_paragraph()
    p_thick2.paragraph_format.space_before = Pt(0)
    p_thick2.paragraph_format.space_after = Pt(0)
    p_thick2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p_thick2.paragraph_format.line_spacing = Pt(1)
    p_thick2.add_run().font.size = Pt(1) 
    add_bottom_border(p_thick2, size='24')
    
    # 3. Reference No & Date
    table3 = doc.add_table(rows=1, cols=2)
    table3.autofit = False
    table3.columns[0].width = Inches(4.5)
    table3.columns[1].width = Inches(2.3)
    
    p_ref = table3.cell(0,0).paragraphs[0]
    p_ref.add_run(f"જા.નં. એસીએન/એન્ટો/{ref_no}/૨૦૨૬, નવસારી")
    
    p_date = table3.cell(0,1).paragraphs[0]
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.add_run(f"તારીખ: {letter_date}")
    
    for cell in table3.rows[0].cells:
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            
    # Add a buffer space before the main letter body
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # 4. Recipient
    p_to = doc.add_paragraph()
    p_to.add_run("પ્રતિ,\n").bold = True
    p_to.add_run("હિસાબ નિયામકશ્રી\nનવસારી કૃષિ યુનિવર્સિટી\nનવસારી- ૩૯૬ ૪૫૦")
    p_to.paragraph_format.space_after = Pt(6)
    
    # 5. Through
    p_through = doc.add_paragraph()
    p_through.add_run("મારફત સવિનય: ").bold = True
    p_through.add_run("આચાર્ય અને ડિનશ્રી , ન. મ. કૃષિ મહાવિદ્યાલય, ન.કૃ.યુ., નવસારી ૩૯૬ ૪૫૦")
    p_through.paragraph_format.space_after = Pt(6)
    
    # 6. Subject
    p_sub = doc.add_paragraph()
    p_sub.add_run("વિષય:- ").bold = True
    p_sub.add_run("બ.સ. ૩૦૩/ ૨૦૯૨ માં ICAR – NCIPM તરફથી આવેલ ગ્રાન્ટ ફાળવવા બાબત...")
    p_sub.paragraph_format.space_after = Pt(12)
    
    # 7. Body Text (With Official Letter Indentation)
    p_body = doc.add_paragraph(body_text)
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_body.paragraph_format.first_line_indent = Inches(0.5) 
    p_body.paragraph_format.space_after = Pt(12)
    
    # 8. Finance Table
    table4 = doc.add_table(rows=2, cols=5)
    table4.style = 'Table Grid'
    headers = ["Name of Centre (Scheme)", "Pay And allowance", "Recurring Contingencies", "Non-Recurring Contingencies", "Total Amount"]
    
    for i, h in enumerate(headers):
        p = table4.cell(0, i).paragraphs[0]
        r_h = p.add_run(h)
        r_h.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    table4.cell(1, 0).text = "AINP on Agril Acarology\n(BH.303/2092)"
    table4.cell(1, 1).text = f"{int(pay_amt):,}/-" if pay_amt > 0 else "-"
    table4.cell(1, 2).text = f"{int(total_rec):,}/-" if total_rec > 0 else "-"
    table4.cell(1, 3).text = f"{int(non_rec_amt):,}/-" if non_rec_amt > 0 else "-"
    table4.cell(1, 4).text = f"{int(total_amt):,}/-"
    
    for i in range(5):
        table4.cell(1, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # 9. Footer Amount & Enclosure
    p_amt = doc.add_paragraph()
    p_amt.add_run(f"In Rupees: {amt_words}").bold = True
    p_amt.paragraph_format.space_after = Pt(12)
    
    p_enc = doc.add_paragraph("સામેલ: ઉપર મુજબ")
    p_enc.paragraph_format.space_after = Pt(24)
    
    # 10. Signatures
    table5 = doc.add_table(rows=1, cols=2)
    table5.autofit = False
    table5.columns[0].width = Inches(3.4)
    table5.columns[1].width = Inches(3.4)
    
    p_sig_left = table5.cell(0,0).paragraphs[0]
    p_sig_left.add_run("પ્રોજેક્ટ ઈન્ચાર્જ").bold = True
    
    p_sig_right = table5.cell(0,1).paragraphs[0]
    p_sig_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sig_right.add_run("પ્રાધ્યાપક અને વડા").bold = True
    
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io
# ---------------------------------------------------------
# 👇 PASTE THIS NEW FUNCTION RIGHT HERE (Around Line 240) 👇
# ---------------------------------------------------------
def create_word_doc(dataframe, month_name, year_num, last_day):
    doc = Document()
    
    # Set narrow margins and change orientation to LANDSCAPE (Horizontal)
    sections = doc.sections
    for section in sections:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
    
    # Headers - Dynamic Month and Year
    title = doc.add_paragraph()
    run_title = title.add_run(f"Statement of Expenditure for the month of {month_name} {year_num}")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title.bold = True
    run_title.font.size = Pt(12)
    
    doc.add_paragraph("Name of the Centre: Navsari").runs[0].bold = True
    doc.add_paragraph("Name of the Scheme: AICRP/AINP on Agricultural Acarology, NAU, Navsari").runs[0].bold = True
    
    # Create Table with 2 header rows
    table = doc.add_table(rows=2, cols=8)
    table.style = 'Table Grid'
    
    # --- ROW 0: Top Level Headers ---
    hdr0 = table.rows[0].cells
    hdr0[0].text = "Sr.\nNo."
    hdr0[1].text = "Head"
    hdr0[2].text = "Opening Balance\nas on 01.04.2025"
    hdr0[3].text = "Funds Received\nfrom the Council\nduring 2025-26"
    
    # Dynamic Headers
    hdr0[4].text = f"Expenditure up to\nthe month of {month_name}\n{year_num}"
    
    # Convert month name back to a digit for the DD.MM.YYYY format
    month_digit = list(calendar.month_name).index(month_name)
    hdr0[5].text = f"Cumulative Expenditure\nup to {last_day}.{month_digit:02d}.{year_num}"
    
    hdr0[7].text = "Total"
    
    # Merge "Cumulative Expenditure" across the two share columns
    hdr0[5].merge(hdr0[6])
    
    # --- ROW 1: Sub Headers ---
    hdr1 = table.rows[1].cells
    hdr1[5].text = "75%\nICAR Share"
    hdr1[6].text = "25%\nState Share"
    
    # Merge vertical columns for headers that span both rows
    for c in [0, 1, 2, 3, 4, 7]:
        table.cell(0, c).merge(table.cell(1, c))
    
    # Format all header cells
    for row_idx in [0, 1]:
        for cell in table.rows[row_idx].cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # --- Add Data Rows ---
    for index, row in dataframe.iterrows():
        row_cells = table.add_row().cells
        is_yellow_header = row.iloc[1] in ["A. Recurring Contingencies", "B. Non Recurring Contingencies"]
        
        for i, cell_data in enumerate(row):
            text_val = str(cell_data) if pd.notna(cell_data) else ""
            row_cells[i].text = text_val
            
            if i == 1:
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            if is_yellow_header or text_val in ["Total - A", "Total - B", "Grand Total A+B"]:
                row_cells[i].paragraphs[0].runs[0].bold = True
                
        # Apply yellow background shading
        if is_yellow_header:
            for cell in row_cells:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'FFFF00')
                tcPr.append(shd)
            
    # Add footer note
    doc.add_paragraph() 
    footer = doc.add_paragraph("In 2025-26 State share released only in Pay and allowances", style='List Bullet')
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
    # ---------------------------------------------------------
# 👇 PASTE THIS NEW FUNCTION RIGHT AFTER `create_word_doc`
# ---------------------------------------------------------
def create_yearly_word_doc(dataframe, fy_string):
    doc = Document()
    
    # Set narrow margins and change orientation to LANDSCAPE
    sections = doc.sections
    for section in sections:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
    
    # Headers
    title = doc.add_paragraph()
    run_title = title.add_run(f"12-Month Expenditure Summary for FY {fy_string}")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title.bold = True
    run_title.font.size = Pt(14)
    
    doc.add_paragraph("Name of the Centre: Navsari").runs[0].bold = True
    doc.add_paragraph("Name of the Scheme: AICRP/AINP on Agricultural Acarology, NAU, Navsari").runs[0].bold = True
    
    columns = dataframe.columns.tolist()
    
    # Create Table
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = 'Table Grid'
    
    # --- Add Column Headers ---
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(columns):
        hdr_cells[i].text = column_name
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # --- Add Data Rows ---
    for index, row in dataframe.iterrows():
        row_cells = table.add_row().cells
        head_val = str(row.iloc[0])
        is_bold_row = head_val in ["A. Recurring Contingencies", "B. Non Recurring Contingencies", "Total - A", "Total - B", "Grand Total A+B"]
        
        for i, cell_data in enumerate(row):
            text_val = str(cell_data) if pd.notna(cell_data) else ""
            row_cells[i].text = text_val
            
            # Left align the Head column, center the numbers
            if i == 0:
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            if is_bold_row:
                row_cells[i].paragraphs[0].runs[0].bold = True
                
        # Apply yellow background shading to Category Rows
        if head_val in ["A. Recurring Contingencies", "B. Non Recurring Contingencies"]:
            for cell in row_cells:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'FFFF00')
                tcPr.append(shd)
            
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
    # ---------------------------------------------------------
# 👇 PASTE THIS NEW FUNCTION RIGHT BELOW `create_word_doc` 👇
# ---------------------------------------------------------
def create_annual_word_doc(dataframe, fy_string):
    doc = Document()
    
    # Set narrow margins and change orientation to LANDSCAPE
    sections = doc.sections
    for section in sections:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
    
    fy_start = fy_string.split('-')[0]
    fy_end = "20" + fy_string.split('-')[1]
    
    # Headers - Dynamic for the Full Year
    title = doc.add_paragraph()
    run_title = title.add_run(f"Statement of Expenditure for the Financial Year {fy_string}")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title.bold = True
    run_title.font.size = Pt(12)
    
    doc.add_paragraph("Name of the Centre: Navsari").runs[0].bold = True
    doc.add_paragraph("Name of the Scheme: AICRP/AINP on Agricultural Acarology, NAU, Navsari").runs[0].bold = True
    
    # Create Table with 2 header rows
    table = doc.add_table(rows=2, cols=8)
    table.style = 'Table Grid'
    
    # --- ROW 0: Top Level Headers ---
    hdr0 = table.rows[0].cells
    hdr0[0].text = "Sr.\nNo."
    hdr0[1].text = "Head"
    hdr0[2].text = f"Opening Balance\nas on 01.04.{fy_start}"
    hdr0[3].text = f"Funds Received\nfrom the Council\nduring {fy_string}"
    hdr0[4].text = f"Expenditure during\nthe FY {fy_string}"
    hdr0[5].text = f"Cumulative Expenditure\nup to 31.03.{fy_end}"
    hdr0[7].text = "Total"
    
    # Merge "Cumulative Expenditure" across the two share columns
    hdr0[5].merge(hdr0[6])
    
    # --- ROW 1: Sub Headers ---
    hdr1 = table.rows[1].cells
    hdr1[5].text = "75%\nICAR Share"
    hdr1[6].text = "25%\nState Share"
    
    # Merge vertical columns for headers that span both rows
    for c in [0, 1, 2, 3, 4, 7]:
        table.cell(0, c).merge(table.cell(1, c))
    
    # Format all header cells
    for row_idx in [0, 1]:
        for cell in table.rows[row_idx].cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # --- Add Data Rows ---
    for index, row in dataframe.iterrows():
        row_cells = table.add_row().cells
        is_yellow_header = row.iloc[1] in ["A. Recurring Contingencies", "B. Non Recurring Contingencies"]
        
        for i, cell_data in enumerate(row):
            text_val = str(cell_data) if pd.notna(cell_data) else ""
            row_cells[i].text = text_val
            
            if i == 1:
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            if is_yellow_header or text_val in ["Total - A", "Total - B", "Grand Total A+B"]:
                row_cells[i].paragraphs[0].runs[0].bold = True
                
        # Apply yellow background shading
        if is_yellow_header:
            for cell in row_cells:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'FFFF00')
                tcPr.append(shd)
            
    # Add footer note
    doc.add_paragraph() 
    footer = doc.add_paragraph("In 2025-26 State share released only in Pay and allowances", style='List Bullet')
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
# ---------------------------------------------------------
# 👇 REPLACE YOUR CURRENT `generate_auc_forwarding_docx` FUNCTION 👇
# ---------------------------------------------------------
def generate_auc_forwarding_docx(ref_no, letter_date, subject_text, body_text):
    """Generates the Forwarding Letter specifically for the AUC."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    style = doc.styles['Normal']
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
        
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(3.6)
    table.columns[2].width = Inches(1.4)
    
    if 'NAU_LOGO' in globals() and NAU_LOGO and os.path.exists(NAU_LOGO):
        cell_left = table.cell(0, 0)
        cell_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_left = cell_left.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_left = p_left.add_run()
        r_left.add_picture(NAU_LOGO, width=Inches(1.8))
        
    p_center = table.cell(0, 1).paragraphs[0]
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_center.paragraph_format.line_spacing = 0.85
    r1 = p_center.add_run("કીટકશાસ્ત્ર વિભાગ\n")
    r1.bold = True
    r1.font.size = Pt(22)
    r2 = p_center.add_run("ન. મ. કૃષિ મહાવિદ્યાલય\nનવસારી કૃષિ યુનિવર્સિટી\nનવસારી- ૩૯૬ ૪૫૦ (ગુજરાત)")
    r2.bold = True
    r2.font.size = Pt(14)
    
    if 'ICAR_LOGO' in globals() and ICAR_LOGO and os.path.exists(ICAR_LOGO):
        p_right = table.cell(0, 2).paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_right = p_right.add_run()
        r_right.add_picture(ICAR_LOGO, width=Inches(1.5))
        
    p_thick1 = doc.add_paragraph()
    p_thick1.paragraph_format.space_before = Pt(0)
    p_thick1.paragraph_format.space_after = Pt(0)
    p_thick1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p_thick1.paragraph_format.line_spacing = Pt(1)
    p_thick1.add_run().font.size = Pt(1) 
    add_bottom_border(p_thick1, size='24')
    
    table2 = doc.add_table(rows=1, cols=2)
    table2.autofit = False
    table2.columns[0].width = Inches(3.4)
    table2.columns[1].width = Inches(3.4)

    p1 = table2.cell(0,0).paragraphs[0]
    p1.add_run("ડૉ. સચિન ડી. પટેલ\nપ્રાધ્યાપક અને વડા")
    p2 = table2.cell(0,1).paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.add_run("મોબાઇલ: +૯૧ ૯૪૨૭૮ ૬૭૯૨૫\nઇમેલ: headentonau@gmail.com")
    
    for cell in table2.rows[0].cells:
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 0.8
            
    p_thick2 = doc.add_paragraph()
    p_thick2.paragraph_format.space_before = Pt(0)
    p_thick2.paragraph_format.space_after = Pt(0)
    p_thick2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p_thick2.paragraph_format.line_spacing = Pt(1)
    p_thick2.add_run().font.size = Pt(1) 
    add_bottom_border(p_thick2, size='24')
    
    # Dynamically grab the year from the date for the ref no
    letter_year = letter_date.split('/')[-1] if '/' in letter_date else "૨૦૨૬"
    
    table3 = doc.add_table(rows=1, cols=2)
    p_ref = table3.cell(0,0).paragraphs[0]
    p_ref.add_run(f"જા.નં. એસીએન/એન્ટો/એઆઈએનપી-એએ/{ref_no}/{letter_year}, નવસારી")
    p_date = table3.cell(0,1).paragraphs[0]
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.add_run(f"તારીખ: {letter_date}")
            
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    p_to = doc.add_paragraph()
    p_to.add_run("પ્રતિ,\n").bold = True
    p_to.add_run("હિસાબ નિયામકશ્રી,\nનવસારી કૃષિ યુનિવર્સિટી,\nનવસારી- ૩૯૬ ૪૫૦.")
    p_to.paragraph_format.space_after = Pt(6)
    
    p_through = doc.add_paragraph()
    p_through.add_run("મારફત સવિનય: ").bold = True
    p_through.add_run("આચાર્ય અને ડીનશ્રી, ન. મ. કૃષિ મહાવિદ્યાલય, ન.કૃ.યુ., નવસારી ૩૯૬ ૪૫૦.")
    p_through.paragraph_format.space_after = Pt(6)
    
    p_sub = doc.add_paragraph()
    p_sub.add_run("વિષય:- ").bold = True
    p_sub.add_run(subject_text)
    p_sub.paragraph_format.space_after = Pt(12)
    
    p_body = doc.add_paragraph(body_text)
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_body.paragraph_format.first_line_indent = Inches(0.5) 
    p_body.paragraph_format.space_after = Pt(24)
    
    p_enc = doc.add_paragraph("સામેલ: ઉપર મુજબ")
    p_enc.paragraph_format.space_after = Pt(36)
    
    table5 = doc.add_table(rows=1, cols=2)
    p_sig_left = table5.cell(0,0).paragraphs[0]
    p_sig_left.add_run("પ્રોજેક્ટ ઈન્ચાર્જ").bold = True
    p_sig_right = table5.cell(0,1).paragraphs[0]
    p_sig_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sig_right.add_run("પ્રાધ્યાપક અને વડા").bold = True
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ---------------------------------------------------------
# 👇 REPLACE YOUR CURRENT `generate_auc_certificate` FUNCTION
# ---------------------------------------------------------
def generate_auc_certificate(inst_data, t1_data, t2_data, cert_text_1, fy_string):
    """Generates the official Audit Utilization Certificate."""
    doc = Document()
    
    # Set Margins
    for section in doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.4)     # <--- REDUCED TOP MARGIN
        section.bottom_margin = Inches(0.4)  # <--- REDUCED BOTTOM MARGIN
    
    # Force Times New Roman Font globally
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
   
    # Certification Text
    p_cert1 = doc.add_paragraph()
    run_cert1 = p_cert1.add_run("     Form of Utilization Certificate & Audit Utilization Certificate")
    run_cert1.bold = True
    run_cert1.font.size = Pt(18)
    run_cert1.font.name = 'Times New Roman'
    
    # Installment Table
    inst_table = doc.add_table(rows=1, cols=3)
    inst_table.style = 'Table Grid'
    hdr = inst_table.rows[0].cells
    hdr[0].text = "Sr.No"
    hdr[1].text = "Letter No and Date"
    hdr[2].text = "Amount"
    
    for row in inst_data:
        cells = inst_table.add_row().cells
        cells[0].text = str(row[0])
        cells[1].text = str(row[1])
        cells[2].text = str(row[2])
        
    for row in inst_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = doc.styles['Normal']
    
    doc.add_paragraph()
        
    # Inject the user-edited text paragraph here
    p_text = doc.add_paragraph(f"{cert_text_1}\n")
    p_text.add_run("2. Certified that I have satisfied myself that the condition on which the expenditure was made have dully fulfilled/are being fulfilled and that I have exercised the following check to see that the money was actually utilized for the purpose for which it was sanctioned.")
    
    p_t1_title = doc.add_paragraph("\nTable 1: Showing the details of receipt and expenditure figure (in Rupees)")
    p_t1_title.runs[0].bold = True
    
    # Table 1
    t1 = doc.add_table(rows=2, cols=5)
    t1.style = 'Table Grid'
    fy_start = fy_string.split('-')[0]
    fy_end = "20" + fy_string.split('-')[1]
    
    t1.cell(0,0).text = f"Opening balance as on 1st April {fy_start}"
    t1.cell(0,1).text = f"Remittance Received {fy_string}"
    t1.cell(0,2).text = "Receipt"
    t1.cell(0,3).text = f"ICAR share of Expenditure during the year {fy_string}"
    t1.cell(0,4).text = f"Closing balance as on 31st March {fy_end}"
    
    for i in range(5):
        t1.cell(1,i).text = str(t1_data[i])
        
    for row in t1.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = doc.styles['Normal']
        
    p_t2_title = doc.add_paragraph("\nTable 2: Showing the head wise details of expenditure figure (in Rupees)")
    p_t2_title.runs[0].bold = True
    
    # Table 2
    t2 = doc.add_table(rows=1, cols=5)
    t2.style = 'Table Grid'
    h2 = t2.rows[0].cells
    h2[0].text = "Head"
    h2[1].text = f"Allocation for the Year {fy_string} (100%)"
    h2[2].text = "ICAR share of Expenditure (75%)"
    h2[3].text = "State Share (25%)"
    h2[4].text = "Total Expenditure"
    
    for row in t2_data:
        cells = t2.add_row().cells
        for i in range(5):
            cells[i].text = str(row[i])
            cells[i].paragraphs[0].style = doc.styles['Normal']
            if str(row[0]) in ["Recurring", "Non Recurring Contingencies", "Total:-"]:
                cells[i].paragraphs[0].runs[0].bold = True

    doc.add_paragraph("\n\n")
    
# Signatures
    sig_table = doc.add_table(rows=1, cols=3)
    sig_table.cell(0,0).text = "PI of the Scheme\nNAU, Navsari"
    sig_table.cell(0,1).text = "Director of Research\nNAU, Navsari"
    sig_table.cell(0,2).text = "Comptroller\nNAU, Navsari"
    
    # Format the top 3 signatures and center them in their columns
    for row in sig_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = doc.styles['Normal']
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add spacing for signatures to be signed
    doc.add_paragraph("\n")
    
    # Add the Chartered Accountant statement centered below the table
    p_ca = doc.add_paragraph("Duly audited and signed by the Chartered Accountant\n\nChartered Accountant")
    p_ca.style = doc.styles['Normal']
    p_ca.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- 3. THE UI APPLICATION ---

def main():
    st.title("🌾 AINP Grant Management System - NAU Navsari")
    
    current_year = datetime.now().year
    fy_options = [f"{y}-{str(y+1)[2:]}" for y in range(current_year-2, current_year+2)]
    selected_fy = st.sidebar.selectbox("Select Financial Year (Apr 1 - Mar 31)", fy_options, index=2)
    
    data = load_data(selected_fy)
    
    tabs = st.tabs([
        "📊 Dashboard", 
        "📤 1. Budget Intake", 
        "💰 2. Installments (PFMS)", 
        "📝 3. Generated Letters", 
        "💸 4. Monthly Spend", 
        "📄 5. SOE Generation", 
        "📜 6. AUC Generation",  # <-- NEW TAB ADDED
        "🤖 AI Chatbot"
    ])
    
    # --- TAB 1: DASHBOARD ---
    with tabs[0]:
        st.header(f"📊 Financial Dashboard (FY {selected_fy})")

        with st.expander("Backup & Audit", expanded=False):
            audit_entries = read_audit_entries(selected_fy, limit=10)
            col_bu1, col_bu2 = st.columns(2)
            with col_bu1:
                backup_buffer = build_fy_backup_zip(selected_fy, data)
                st.download_button(
                    "Download FY Backup ZIP",
                    data=backup_buffer,
                    file_name=f"AINP_Backup_{selected_fy}.zip",
                    mime="application/zip",
                    key=f"backup_zip_{selected_fy}",
                )
            with col_bu2:
                audit_path = get_audit_filename(selected_fy)
                audit_bytes = load_document_bytes(audit_path) if os.path.exists(audit_path) else b""
                st.download_button(
                    "Download Audit Log",
                    data=audit_bytes,
                    file_name=f"audit_log_{selected_fy}.jsonl",
                    mime="application/jsonl",
                    key=f"audit_log_download_{selected_fy}",
                    disabled=not bool(audit_bytes),
                )

            if audit_entries:
                st.write("Recent changes")
                st.dataframe(pd.DataFrame(audit_entries), width="stretch", hide_index=True)
            else:
                st.info("No audit entries recorded yet.")

            restore_file = st.file_uploader("Restore FY JSON backup", type=["json"], key=f"restore_json_{selected_fy}")
            if restore_file:
                restore_data, restore_error = preview_restore_json(restore_file, selected_fy)
                if restore_error:
                    st.error(f"Restore file is not valid: {restore_error}")
                else:
                    st.success("Restore file is valid. Preview:")
                    st.json({
                        "financial_year": restore_data.get("financial_year"),
                        "installments": len(restore_data.get("installments", [])),
                        "expenditure": len(restore_data.get("expenditure", [])),
                        "has_allocation": bool(restore_data.get("allocation")),
                        "has_revised_allocation": bool(restore_data.get("revised_allocation")),
                    })
                    confirm_restore = st.checkbox("I understand this will replace the current FY data.", key=f"confirm_restore_{selected_fy}")
                    if st.button("Restore This FY JSON", key=f"restore_btn_{selected_fy}", disabled=not confirm_restore):
                        save_data(
                            restore_data,
                            selected_fy,
                            audit_action="restore_fy_json",
                            audit_details={"source_file": restore_file.name},
                        )
                        st.success("FY data restored.")
                        st.rerun()
        
        # --- 1. DATA PREPARATION ---
        # Get Active Allocation
        active_budget = data.get('revised_allocation') if data.get('revised_allocation') else data.get('allocation', {})
        total_allocated = sum(v.get('total', 0) for v in active_budget.values()) if active_budget else 0.0
        
        # Get Total Received
        total_received = sum(coerce_amount(inst.get('amount')) for inst in data.get('installments', []))
        
        # Get Total Spent
        df_exp = pd.DataFrame(data.get('expenditure', []))
        total_spent = df_exp['amount'].sum() if not df_exp.empty else 0.0
        
        # Calculate Available Physical Balance
        available_balance = total_received - total_spent
        
        # Quick Matcher for Heads (for received funds)
        def dash_match_head(raw_string):
            rs = str(raw_string).upper()
            if "PAY" in rs or "ESTABLISHMENT" in rs: return "Pay and Allowances"
            if "TA" in rs or "TRAVELLING" in rs: return "Travelling Allowances (TA)"
            if "TSP" in rs: return "TSP"
            if "NON" in rs or "EQUIP" in rs or "WORK" in rs: return "Non-Recurring Contingencies (Equipments/Works)" 
            if "ORC" in rs or "CONTINGENC" in rs or "RECURRING" in rs: return "Other Recurring Contingencies (ORC)"
            return None

        # Build Head-wise Data
        head_stats = []
        for head in BUDGET_HEADS:
            alloc = active_budget.get(head, {}).get('total', 0.0)
            
            recv = 0.0
            for inst in data.get('installments', []):
                for h, amt in inst.get('heads', {}).items():
                    if dash_match_head(h) == head:
                        recv += float(amt)
                        
            spent = 0.0
            if not df_exp.empty:
                spent = df_exp[df_exp['head'] == head]['amount'].sum()
                
            head_stats.append({
                "Budget Head": head,
                "Allocated (₹)": alloc,
                "Received (₹)": recv,
                "Spent (₹)": spent,
                "Remaining (Received - Spent)": recv - spent
            })
            
        df_heads = pd.DataFrame(head_stats).set_index("Budget Head")

        # --- 2. TOP METRICS ROW ---
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Sanctioned Budget", f"₹{total_allocated:,.0f}")
        with col2:
            recv_pct = (total_received / total_allocated * 100) if total_allocated > 0 else 0
            st.metric("Total Funds Received", f"₹{total_received:,.0f}", f"{recv_pct:.1f}% of Sanctioned", delta_color="normal")
        with col3:
            spent_pct = (total_spent / total_received * 100) if total_received > 0 else 0
            st.metric("Total Expenditure", f"₹{total_spent:,.0f}", f"{spent_pct:.1f}% of Received utilized", delta_color="inverse")
        with col4:
            st.metric("Available Cash Balance", f"₹{available_balance:,.0f}", f"₹{total_allocated - total_received:,.0f} Pending from ICAR", delta_color="off")

        st.divider()

        # --- 3. CONCERNS & ALERTS SECTION ---
        alerts = []
        if total_allocated > total_received:
            alerts.append(f"⏳ **Pending Grant:** ₹{total_allocated - total_received:,.2f} is sanctioned but not yet received from the Council.")
            
        for idx, row in df_heads.iterrows():
            if row['Spent (₹)'] > row['Received (₹)']:
                alerts.append(f"🚨 **Deficit Alert:** Expenditure in **{idx}** (₹{row['Spent (₹)']:,.0f}) exceeds the funds currently received (₹{row['Received (₹)']:,.0f}) for this head!")
            elif row['Allocated (₹)'] > 0 and (row['Spent (₹)'] / row['Allocated (₹)']) > 0.9:
                alerts.append(f"⚠️ **Near Empty:** You have utilized over 90% of the total sanctioned budget for **{idx}**.")

        if alerts:
            with st.expander("🚨 Actionable Alerts & Financial Concerns (Click to expand)", expanded=True):
                for alert in alerts:
                    if "🚨" in alert: st.error(alert)
                    elif "⚠️" in alert: st.warning(alert)
                    else: st.info(alert)
        else:
            st.success("✅ Financial Health is stable. No immediate concerns or deficits detected.")

        # --- 4. VISUALIZATIONS ---
        col_chart, col_table = st.columns([1.5, 1])
        
        with col_chart:
            st.subheader("Head-wise Financial Status")
            # Bar chart comparing Allocated vs Received vs Spent
            st.bar_chart(df_heads[["Allocated (₹)", "Received (₹)", "Spent (₹)"]], height=350)
            
        with col_table:
            st.subheader("Head-wise Details")
            # Display clean table with numbers
            st.dataframe(df_heads.style.format("{:,.0f}"), width="stretch", height=350)

        # --- 5. MONTHLY BURN RATE (TREND) ---
        st.divider()
        st.subheader("📈 Monthly Expenditure Trend")
        if not df_exp.empty:
            df_exp['date'] = pd.to_datetime(df_exp['date'], errors="coerce")
            df_exp = df_exp.dropna(subset=['date'])
            # Extract Year-Month for grouping (e.g., "2025-04")
            df_exp['Month'] = df_exp['date'].dt.to_period('M').astype(str)
            monthly_trend = df_exp.groupby('Month')['amount'].sum().reset_index()
            monthly_trend = monthly_trend.set_index('Month')
            
            st.bar_chart(monthly_trend, height=300, color="#FF4B4B")
        else:
            st.info("No expenditures logged yet to display trends. Go to Tab 4 to add expenses.")

    # --- TAB 2: BUDGET INTAKE ---
    with tabs[1]:
        st.header("Upload Budget Allocation / Revision PDF")
        
        st.markdown("""
        **Instructions for Uploading:**
        Please describe the document in the text box below (e.g., 'Initial Allocation', 'Revised Allocation', 'Q1 Release', 'Q3 Release') and upload the corresponding PDF.
        """)
        
        doc_type_input = st.text_input("Describe the document you are uploading:")
        budget_file = st.file_uploader("Upload Budget Document", type=['pdf'], key="budget_up")
        
        if budget_file and st.button("Analyze & Process Budget"):
            if not doc_type_input.strip():
                st.warning("⚠️ Please provide a description in the text box above before processing.")
            else:
                with st.spinner(f"AI is analyzing the PDF..."):
                    
                    budget_prompt = {
                        "context": f"Document Description: {doc_type_input}. Scheme: AICRP/AINP Acarology. Financial Year: {selected_fy}. Extract the budget allocation table. IMPORTANT: Map the head names exactly to: 'Pay and Allowances', 'Travelling Allowances (TA)', 'Other Recurring Contingencies (ORC)', 'Non-Recurring Contingencies (Equipments/Works)', 'TSP'. Convert all Lakh values to absolute Rupees (e.g., 1.33 Lakhs becomes 133000, 40 Lakhs becomes 4000000).",
                        "structure": {
                            "is_revision": "boolean (true if this revises a previous allocation)",
                            "date": "YYYY-MM-DD",
                            "heads": [
                                {"head_name": "string", "icar_share": 0.0, "state_share": 0.0, "total": 0.0}
                            ]
                        }
                    }
                    
                    extracted_budget = process_upload_with_ai(budget_file, budget_prompt)
                    
                    if extracted_budget:
                        st.success("Document analyzed successfully!")
                        
                        # Force creation of a dictionary with ALL 5 standard heads (defaulted to 0)
                        budget_dict = {h: {'icar': 0.0, 'state': 0.0, 'total': 0.0} for h in BUDGET_HEADS}
                        
                        # --- STRICT 75:25 MATHEMATICAL ENFORCEMENT (WITH 100% TSP EXCEPTION) ---
                        for head in extracted_budget.get('heads', []):
                            extracted_total = float(head.get('total') or 0.0)
                            extracted_icar = float(head.get('icar_share') or 0.0)
                            raw_name = head.get('head_name', '').upper()
                            
                            # Match the AI's extracted name to the standard BUDGET_HEADS
                            matched_head = None
                            if "PAY" in raw_name: matched_head = BUDGET_HEADS[0]
                            elif "TRAV" in raw_name or "TA" in raw_name: matched_head = BUDGET_HEADS[1]
                            elif "NON" in raw_name or "EQUIP" in raw_name or "WORK" in raw_name: matched_head = BUDGET_HEADS[3]
                            elif "RECURR" in raw_name or "ORC" in raw_name: matched_head = BUDGET_HEADS[2]
                            elif "TSP" in raw_name: matched_head = BUDGET_HEADS[4]
                            
                            if matched_head:
                                # Exception: TSP is always 100% ICAR share
                                if "TSP" in matched_head.upper():
                                    final_total = extracted_total if extracted_total > 0 else extracted_icar
                                    budget_dict[matched_head]['total'] = round(final_total, 2)
                                    budget_dict[matched_head]['icar'] = round(final_total, 2)
                                    budget_dict[matched_head]['state'] = 0.0
                                else:
                                    # Standard 75:25 Split
                                    if extracted_total == 0 and extracted_icar > 0:
                                        final_total = extracted_icar / 0.75
                                    else:
                                        final_total = extracted_total

                                    budget_dict[matched_head]['total'] = round(final_total, 2)
                                    budget_dict[matched_head]['icar'] = round(final_total * 0.75, 2)
                                    budget_dict[matched_head]['state'] = round(final_total * 0.25, 2)
                        
                        doc_date = extracted_budget.get('date', 'Unknown')
                        desc_lower = doc_type_input.lower()
                        assigned_type = "Unknown"
                        dict_key = ""
                        
                        if "revis" in desc_lower:
                            data['revised_allocation'] = budget_dict
                            assigned_type = "Revised Allocation"
                            dict_key = "revised_allocation"
                        elif "q1" in desc_lower or "1st" in desc_lower:
                            data['quarterly_allocations']["Q1"] = budget_dict
                            assigned_type = "Q1 Release"
                            dict_key = "Q1"
                        elif "q2" in desc_lower or "2nd" in desc_lower:
                            data['quarterly_allocations']["Q2"] = budget_dict
                            assigned_type = "Q2 Release"
                            dict_key = "Q2"
                        elif "q3" in desc_lower or "3rd" in desc_lower:
                            data['quarterly_allocations']["Q3"] = budget_dict
                            assigned_type = "Q3 Release"
                            dict_key = "Q3"
                        elif "q4" in desc_lower or "4th" in desc_lower:
                            data['quarterly_allocations']["Q4"] = budget_dict
                            assigned_type = "Q4 Release"
                            dict_key = "Q4"
                        else:
                            data['allocation'] = budget_dict
                            assigned_type = "Initial Allocation"
                            dict_key = "allocation"
                            
                        # Save the PDF locally and to GitHub for later download
                        pdf_path = save_uploaded_document(budget_file, f"{selected_fy}_{dict_key}.pdf")
                            
                        st.info(f"📅 **Document Date:** {doc_date} | 🔄 **Detected Type:** {assigned_type}")
                        
                        data['latest_quarter'] = assigned_type
                        data['latest_date'] = doc_date
                        
                        save_data(
                            data,
                            selected_fy,
                            audit_action="budget_upload",
                            audit_details={
                                "document_type": assigned_type,
                                "document_date": doc_date,
                                "file": os.path.basename(pdf_path),
                            },
                        )
                        st.toast("Budget Data Saved to GitHub & PDF saved locally!")

        # --- BUDGET VISUALIZATION & EDITING ---
        st.divider()
        st.subheader("📑 Uploaded Budget Documents (Editable)")
        
        # Interactive budget editor function
        def interactive_budget_editor(b_dict, dict_key, doc_type_name):
            df_data = []
            for h in BUDGET_HEADS:
                if h in b_dict:
                    df_data.append({
                        "Budget Head": h,
                        "ICAR Share (₹)": b_dict[h]['icar'],
                        "State Share (₹)": b_dict[h]['state'],
                        "Total (₹)": b_dict[h]['total']
                    })
            
            if not df_data:
                return

            df = pd.DataFrame(df_data)
            
            # Display interactive Data Editor
            edited_df = st.data_editor(df, width="stretch", hide_index=True, key=f"editor_{dict_key}")
            
            # Read-only totals for display below the editor
            tot_icar = edited_df["ICAR Share (₹)"].astype(float).sum()
            tot_state = edited_df["State Share (₹)"].astype(float).sum()
            tot_all = edited_df["Total (₹)"].astype(float).sum()
            st.markdown(f"**Calculated Totals ➡️ ICAR:** ₹{tot_icar:,.2f} | **State:** ₹{tot_state:,.2f} | **Grand Total:** ₹{tot_all:,.2f}")
            
            col1, col2 = st.columns([1, 3])
            with col1:
                if st.button(f"💾 Save & Recalculate", key=f"save_{dict_key}"):
                    new_dict = {}
                    for index, row in edited_df.iterrows():
                        head = row["Budget Head"]
                        raw_total = float(row["Total (₹)"])
                        raw_icar = float(row["ICAR Share (₹)"])
                        
                        if "TSP" in head.upper():
                            final_total = raw_total if raw_total > 0 else raw_icar
                            new_dict[head] = {'icar': final_total, 'state': 0.0, 'total': final_total}
                        else:
                            if raw_total == 0 and raw_icar > 0:
                                final_total = raw_icar / 0.75
                            else:
                                final_total = raw_total

                            new_dict[head] = {
                                'icar': round(final_total * 0.75, 2),
                                'state': round(final_total * 0.25, 2),
                                'total': round(final_total, 2)
                            }
                    
                    if dict_key == 'allocation':
                        data['allocation'] = new_dict
                    elif dict_key == 'revised_allocation':
                        data['revised_allocation'] = new_dict
                    else:
                        data['quarterly_allocations'][dict_key] = new_dict
                        
                    save_data(
                        data,
                        selected_fy,
                        audit_action="budget_edit",
                        audit_details={"document_type": doc_type_name, "key": dict_key},
                    )
                    st.success("Changes Saved & Recalculated!")
                    st.rerun()

            with col2:
                pdf_path = safe_documents_path(f"{selected_fy}_{dict_key}.pdf")
                pdf_bytes = load_document_bytes(pdf_path)
                if pdf_bytes:
                    st.download_button(
                        label=f"📥 Download Uploaded PDF",
                        data=pdf_bytes,
                        file_name=f"{doc_type_name}_{selected_fy}.pdf",
                        mime="application/pdf",
                        key=f"dl_{dict_key}",
                    )

        col_a, col_b = st.columns(2)
        with col_a:
            if data.get('allocation'):
                with st.expander("📄 Initial Full-Year Budget Allocation"):
                    interactive_budget_editor(data['allocation'], 'allocation', "Initial_Allocation")
        with col_b:
            if data.get('revised_allocation'):
                with st.expander("📄 Revised Full-Year Budget Allocation"):
                    interactive_budget_editor(data['revised_allocation'], 'revised_allocation', "Revised_Allocation")
                    
        st.write("📅 **Quarterly Releases (Click to edit data)**")
        q_cols = st.columns(4)
        for i, q in enumerate(["Q1", "Q2", "Q3", "Q4"]):
            q_data = data['quarterly_allocations'].get(q)
            if q_data:
                with q_cols[i]:
                    with st.expander(f"📄 {q} Release Data"):
                        interactive_budget_editor(q_data, q, f"{q}_Release")
                        
        st.divider()
        
        # --- Total Mismatch/Status Table ---
        active_budget = data.get('revised_allocation') or data.get('allocation')
        if active_budget:
            st.subheader("⚖️ Allocation vs. Quarterly Release Mismatch")
            mismatch_data = {}
            
            # Force standard order
            for head in BUDGET_HEADS:
                vals = active_budget.get(head, {})
                tot_alloc = vals.get('total', 0)
                
                # Sum the quarterly totals for this specific head
                q_sum = 0
                for q in ["Q1", "Q2", "Q3", "Q4"]:
                    if data['quarterly_allocations'].get(q):
                        q_sum += data['quarterly_allocations'][q].get(head, {}).get('total', 0)
                        
                mismatch = tot_alloc - q_sum
                
                # Only show rows if there is allocated budget OR if money was mysteriously released for it
                if tot_alloc > 0 or q_sum > 0:
                    mismatch_data[head] = {
                        "Total Sanctioned (₹)": round(tot_alloc, 2),
                        "Released Q1-Q4 (₹)": round(q_sum, 2),
                        "Pending/Mismatch (₹)": round(mismatch, 2)
                    }
            
            if mismatch_data:
                df_mismatch = pd.DataFrame.from_dict(mismatch_data, orient='index')
                df_mismatch.loc['TOTAL'] = df_mismatch.sum(numeric_only=True)
                st.dataframe(df_mismatch, width="stretch")
            else:
                st.info("No allocated budget values to show yet.")
        else:
            st.info("Upload budget allocation to view Mismatch Table.")

    # --- TAB 3: INSTALLMENTS ---
    with tabs[2]:
        st.header("Received Grant Installment (PFMS & Email)")
        st.write("Upload the official release email and the PFMS attachment to log funds.")
        
        col1, col2 = st.columns(2)
        with col1:
            email_file = st.file_uploader("1. Upload Email PDF", type=['pdf'], key="email_up")
        with col2:
            pfms_file = st.file_uploader("2. Upload PFMS Receipt PDF", type=['pdf'], key="pfms_up")

        if email_file and pfms_file and st.button("Analyze & Process Installment"):
            with st.spinner("Analyzing Email and PFMS documents..."):
                
                # Package BOTH PDFs as native files for Gemini
                pdf_data_email = types.Part.from_bytes(data=email_file.getvalue(), mime_type="application/pdf")
                pdf_data_pfms = types.Part.from_bytes(data=pfms_file.getvalue(), mime_type="application/pdf")
                
                full_prompt = f"""
                Analyze the provided Email and PFMS documents and extract the installment information.
                Map the monetary values exactly to these heads: 'Pay and Allowances', 'Travelling Allowances (TA)', 'Other Recurring Contingencies (ORC)', 'Non-Recurring Contingencies (Equipments/Works)', 'TSP'.
                
                REQUIRED JSON STRUCTURE:
                {{
                    "date": "YYYY-MM-DD",
                    "installment_number": "String (e.g., I, II, III, IV, V, etc.)",
                    "purpose": "String (e.g., GIA-Salary, GIA-General, GIA-TSP, GIA-General & Capital)",
                    "pfms_transaction_id": "String (e.g., C022...)",
                    "heads": [
                        {{"head_name": "string", "amount": 0.0}}
                    ]
                }}
                IMPORTANT: Return ONLY valid JSON.
                """
                
                extracted_inst = extract_structured_with_ai(
                    [pdf_data_email, pdf_data_pfms],
                    full_prompt,
                    InstallmentExtraction,
                )
                if extracted_inst:
                    st.session_state['pending_installment'] = extracted_inst
                    st.success("Documents analyzed successfully!")

        # If we have a pending installment in session state, show the editor
        if 'pending_installment' in st.session_state:
            extracted = st.session_state['pending_installment']
            
            dt_obj = parse_date(extracted.get('date')) or datetime.combine(date.today(), datetime.min.time())
                
            st.write("✏️ **Step 1: Edit Installment Details (Fix any AI extraction errors here):**")
            mc1, mc2, mc3, mc4 = st.columns(4)
            edit_date = mc1.date_input("Date", dt_obj)
            edit_inst_no = mc2.text_input("Inst No. (e.g. I, II)", extracted.get('installment_number', ''))
            edit_purpose = mc3.text_input("Purpose", extracted.get('purpose', ''))
            edit_pfms_id = mc4.text_input("PFMS ID", extracted.get('pfms_transaction_id', ''))
            
            # Recalculate quarter based on edited date
            m = edit_date.month
            if m in [4, 5, 6]: q_str = "Q1"
            elif m in [7, 8, 9]: q_str = "Q2"
            elif m in [10, 11, 12]: q_str = "Q3"
            else: q_str = "Q4"
            
            # Prepare dataframe for editing
            inst_dict = {h: 0.0 for h in BUDGET_HEADS}
            for h in extracted.get('heads', []):
                raw_name = h.get('head_name', '').upper()
                amt = float(h.get('amount') or 0.0)
                matched_head = None
                if "PAY" in raw_name or "SALARY" in raw_name: matched_head = BUDGET_HEADS[0]
                elif "TRAV" in raw_name or "TA" in raw_name: matched_head = BUDGET_HEADS[1]
                elif "NON" in raw_name or "EQUIP" in raw_name or "WORK" in raw_name or "CAPITAL" in raw_name: matched_head = BUDGET_HEADS[3]
                elif "RECURR" in raw_name or "ORC" in raw_name or "GENERAL" in raw_name: matched_head = BUDGET_HEADS[2]
                elif "TSP" in raw_name: matched_head = BUDGET_HEADS[4]
                
                if matched_head:
                    inst_dict[matched_head] += amt
                    
            df_data = [{"Budget Head": k, "Amount (₹)": v} for k, v in inst_dict.items()]
            df = pd.DataFrame(df_data)
            
            st.write("✏️ **Step 2: Review and Edit Installment Amounts:**")
            edited_df = st.data_editor(df, width="stretch", hide_index=True, key="inst_editor")
            
            total_amt = edited_df["Amount (₹)"].astype(float).sum()
            st.markdown(f"**Total Installment Amount:** ₹{total_amt:,.2f}")
            
            col_save, col_cancel = st.columns([1, 4])
            with col_save:
                if st.button("💾 Save Installment"):
                    final_heads = {row["Budget Head"]: float(row["Amount (₹)"]) for _, row in edited_df.iterrows()}
                    
                    new_inst = {
                        "date": edit_date.strftime("%Y-%m-%d"),
                        "quarter": q_str,
                        "installment_num": edit_inst_no,
                        "purpose": edit_purpose,
                        "pfms_id": edit_pfms_id,
                        "amount": total_amt, 
                        "heads": final_heads,
                        "type": edit_inst_no, 
                        "available": False
                    }
                    
                    safe_pfms_id = sanitize_filename(new_inst['pfms_id'], "UNKNOWN_PFMS")
                    if not new_inst['pfms_id'].strip():
                        st.warning("Please enter a PFMS ID before saving.")
                    elif not any(inst.get('pfms_id') == new_inst['pfms_id'] for inst in data['installments']):
                        data['installments'].append(new_inst)
                        save_data(
                            data,
                            selected_fy,
                            audit_action="installment_add",
                            audit_details={"pfms_id": new_inst['pfms_id'], "amount": total_amt},
                        )
                        
                        # Save PDFs locally
                        save_uploaded_document(email_file, f"{selected_fy}_Inst_{safe_pfms_id}_Email.pdf")
                        save_uploaded_document(pfms_file, f"{selected_fy}_Inst_{safe_pfms_id}_PFMS.pdf")
                        
                        st.toast("Installment Saved and Backed up!")
                        del st.session_state['pending_installment']
                        st.rerun()
                    else:
                        st.warning("An installment with this PFMS ID already exists.")
            with col_cancel:
                if st.button("❌ Cancel"):
                    del st.session_state['pending_installment']
                    st.rerun()
                    
       # Display existing installments
        st.divider()
        st.subheader("📁 Saved Installments (Editable)")
        if data['installments']:
            # Sort by Date
            data['installments'] = sorted(data['installments'], key=lambda x: x.get('date', '2000-01-01'))
            
            for q in ["Q1", "Q2", "Q3", "Q4"]:
                q_insts = [i for i in data['installments'] if i.get('quarter') == q]
                if q_insts:
                    with st.expander(f"📦 {q} Installments ({len(q_insts)})"):
                        for inst in q_insts:
                            # Allow editing basic info
                            st.write(f"**Date:** {inst.get('date')} | **PFMS ID:** {inst.get('pfms_id')}")
                            
                            mc1, mc2 = st.columns(2)
                            new_inst_num = mc1.text_input("Inst No:", inst.get('installment_num', ''), key=f"enum_{inst['pfms_id']}")
                            new_purpose = mc2.text_input("Purpose:", inst.get('purpose', ''), key=f"epur_{inst['pfms_id']}")
                            
                            col_a, col_b = st.columns([3, 1])
                            with col_a:
                                # Provide an editable table for the saved installment
                                inst_heads = inst.get('heads', {h: 0.0 for h in BUDGET_HEADS})
                                df_data = [{"Budget Head": k, "Amount (₹)": v} for k, v in inst_heads.items()]
                                df_saved = pd.DataFrame(df_data)
                                
                                edited_saved_df = st.data_editor(df_saved, width="stretch", hide_index=True, key=f"edit_saved_{inst['pfms_id']}")
                                tot_amt_saved = edited_saved_df["Amount (₹)"].astype(float).sum()
                                st.markdown(f"**Total Amount:** ₹{tot_amt_saved:,.2f}")
                                
                                c_save, c_del = st.columns(2)
                                with c_save:
                                    if st.button("💾 Save Changes", key=f"save_btn_{inst['pfms_id']}"):
                                        final_heads_saved = {row["Budget Head"]: float(row["Amount (₹)"]) for _, row in edited_saved_df.iterrows()}
                                        for main_inst in data['installments']:
                                            if main_inst['pfms_id'] == inst['pfms_id']:
                                                main_inst['heads'] = final_heads_saved
                                                main_inst['amount'] = tot_amt_saved
                                                main_inst['installment_num'] = new_inst_num
                                                main_inst['purpose'] = new_purpose
                                                main_inst['type'] = new_inst_num if "State Share" not in inst.get('type','') else new_inst_num + " (State Share)"
                                                break
                                        save_data(
                                            data,
                                            selected_fy,
                                            audit_action="installment_edit",
                                            audit_details={"pfms_id": inst['pfms_id'], "amount": tot_amt_saved},
                                        )
                                        st.toast("Installment Updated!")
                                        st.rerun()
                                with c_del:
                                    if st.button("🗑️ Delete Installment", key=f"del_btn_{inst['pfms_id']}"):
                                        data['installments'] = [i for i in data['installments'] if i['pfms_id'] != inst['pfms_id']]
                                        save_data(
                                            data,
                                            selected_fy,
                                            audit_action="installment_delete",
                                            audit_details={"pfms_id": inst['pfms_id']},
                                        )
                                        st.toast("Installment Deleted!")
                                        st.rerun()
                                        
                            with col_b:
                                safe_pfms_id = sanitize_filename(inst['pfms_id'], "UNKNOWN_PFMS")
                                email_path = safe_documents_path(f"{selected_fy}_Inst_{safe_pfms_id}_Email.pdf")
                                pfms_path = safe_documents_path(f"{selected_fy}_Inst_{safe_pfms_id}_PFMS.pdf")
                                email_bytes = load_document_bytes(email_path)
                                pfms_bytes = load_document_bytes(pfms_path)
                                if email_bytes:
                                    st.download_button("📥 Download Email", email_bytes, file_name=f"{safe_pfms_id}_Email.pdf", key=f"dl_e_{inst['pfms_id']}")
                                if pfms_bytes:
                                    st.download_button("📥 Download PFMS", pfms_bytes, file_name=f"{safe_pfms_id}_PFMS.pdf", key=f"dl_p_{inst['pfms_id']}")
                            st.divider()
                               
            
            # --- SUMMARY TABLE ---
            st.divider()
            st.subheader("📊 Summary of Received Installments (Q1 - Q4)")
            
            summary_data = {h: {"Q1 (₹)": 0.0, "Q2 (₹)": 0.0, "Q3 (₹)": 0.0, "Q4 (₹)": 0.0, "Total (₹)": 0.0} for h in BUDGET_HEADS}
            
            for inst in data['installments']:
                q = inst.get('quarter')
                if q in ["Q1", "Q2", "Q3", "Q4"]:
                    q_key = f"{q} (₹)"
                    for head, amt in inst.get('heads', {}).items():
                        if head in summary_data:
                            summary_data[head][q_key] += float(amt)
                            summary_data[head]["Total (₹)"] += float(amt)
                            
            df_summary = pd.DataFrame.from_dict(summary_data, orient='index')
            df_summary.loc['GRAND TOTAL'] = df_summary.sum(numeric_only=True)
            st.dataframe(df_summary, width="stretch")

        else:
            st.info("No installments recorded yet.")

    # --- TAB 4: GENERATED LETTERS ---
    with tabs[3]:
        st.header("Draft Letters based on PFMS Receipts")
        st.write("Generate the Gujarati letter (Word Document) to send to Comptroller.")
        
        pending_utilization = [inst for inst in data['installments'] if not inst.get('utilization_letter_generated')]
        
        if pending_utilization:
            options = {}
            for inst in pending_utilization:
                pfms_id = inst.get('pfms_id') or "UNKNOWN_PFMS"
                label = f"{inst.get('type', 'Installment')} (₹{coerce_amount(inst.get('amount')):,.2f} - {pfms_id})"
                options[label] = inst
            selected_inst_str = st.selectbox("Select PFMS Receipt to draft letter for:", list(options.keys()))
            selected_inst_data = options[selected_inst_str]

            # Extract heads for the template
            inst_heads = selected_inst_data.get('heads', {})
            pay_amt = inst_heads.get('Pay and Allowances', 0)
            rec_amt = inst_heads.get('Other Recurring Contingencies (ORC)', 0)
            ta_amt = inst_heads.get('Travelling Allowances (TA)', 0)
            total_rec = rec_amt + ta_amt
            non_rec_amt = inst_heads.get('Non-Recurring Contingencies (Equipments/Works)', 0)
            total_amt = selected_inst_data.get('amount', 0)

            # Editor Fields
            st.subheader("✏️ Edit Letter Details")
            col_a, col_b = st.columns(2)
            with col_a:
                ref_no = st.text_input("Reference No. (જા.નં. એસીએન/એન્ટો/___/૨૦૨૬):", value="")
                letter_date = st.text_input("Date (તારીખ):", value=datetime.now().strftime("%d/%m/%Y"))
            with col_b:
                amt_words = st.text_input("Amount in Words (In Rupees):", value="One lakh rupee only")
                
            body_text = st.text_area("Body Text:", value="જય ભારત સહ ઉપરોક્ત વિષય અન્વયે જણાવવાનું કે, અત્રેના કીટકશાસ્ત્ર વિભાગ ખાતે ચાલતી આઈ.સી.એ.આર. યોજના AINP on Agricultural Acarology (75:25%) (BH.303/2092) માં આવેલ ગ્રાન્ટને કોષ્ટકમાં જણાવ્યાનુસાર ફાળવી આપવા આપ સાહેબશ્રીને નમ્ર વિનંતી.", height=100)

            st.divider()
            st.subheader("👀 Letter Preview")

            # Inline CSS matching the requested exact layout
            st.markdown("""
            <style>
                .block-container { padding-top: 2rem; padding-bottom: 2rem; background-color: #ffffff; }
                .letter-body { font-family: 'Arial', sans-serif; font-size: 16px; color: #000000; line-height: 1.5; }
                .bold { font-weight: bold; }
                .center { text-align: center; }
                .right { text-align: right; }
                .justify { text-align: justify; }
                .indent { text-indent: 50px; }
                .table-custom { width: 100%; border-collapse: collapse; margin: 20px 0; font-size: 15px; }
                .table-custom th, .table-custom td { border: 1px solid black; padding: 8px; text-align: center; color: black; }
            </style>
            """, unsafe_allow_html=True)

            col1, col2, col3 = st.columns([1.2, 3, 1.2])
            with col1:
                if NAU_LOGO and os.path.exists(NAU_LOGO):
                    st.image(NAU_LOGO, width="stretch")
            with col2:
                st.markdown("""
                <div class="letter-body center">
                    <span class="bold">કીટકશાસ્ત્ર વિભાગ</span><br>
                    ન. મ. કૃષિ મહાવિદ્યાલય<br>
                    નવસારી કૃષિ યુનિવર્સિટી<br>
                    નવસારી- ૩૯૬ ૪૫૦ (ગુજરાત)
                </div>
                """, unsafe_allow_html=True)
            with col3:
                if ICAR_LOGO and os.path.exists(ICAR_LOGO):
                    st.image(ICAR_LOGO, width="stretch")

            st.markdown("<hr style='border: 1px solid black; margin-top: 10px; margin-bottom: 15px;' />", unsafe_allow_html=True)
            
            # Safely format numbers with commas
            pay_val = f"{int(pay_amt):,}/-" if pay_amt > 0 else "-"
            rec_val = f"{int(total_rec):,}/-" if total_rec > 0 else "-"
            non_rec_val = f"{int(non_rec_amt):,}/-" if non_rec_amt > 0 else "-"
            tot_val = f"{int(total_amt):,}/-"

            st.markdown(f"""
            <div class="letter-body">
                <div style="display: flex; justify-content: space-between;">
                    <div><span class="bold">ડૉ. જે. જે. પસ્તાગિયા</span><br>પ્રાધ્યાપક અને વડા (ઈ/ચા.)</div>
                    <div class="right">મોબાઇલ: +૯૧ ૯૮૭૯૦ ૩૮૫૩૯<br>ઇમેલ: headentonau@gmail.com</div>
                </div>
                <br>
                <div style="display: flex; justify-content: space-between;">
                    <div>જા.નં. એસીએન/એન્ટો/{ref_no}/૨૦૨૬, નવસારી</div>
                    <div class="right">તારીખ: {letter_date}</div>
                </div>
                <br>
                <div><span class="bold">પ્રતિ,</span><br>હિસાબ નિયામકશ્રી<br>નવસારી કૃષિ યુનિવર્સિટી<br>નવસારી- ૩૯૬ ૪૫૦</div>
                <br>
                <div><span class="bold">મારફત સવિનય:</span> આચાર્ય અને ડિનશ્રી , ન. મ. કૃષિ મહાવિદ્યાલય, ન.કૃ.યુ., નવસારી ૩૯૬ ૪૫૦</div>
                <br>
                <div><span class="bold">વિષય:-</span> બ.સ. ૩૦૩/ ૨૦૯૨ માં ICAR – NCIPM તરફથી આવેલ ગ્રાન્ટ ફાળવવા બાબત...</div>
                <br>
                <div class="justify indent">{body_text.replace(chr(10), '<br>')}</div>
                <table class="table-custom">
                    <tr>
                        <th style="text-align: left;">Name of Centre (Scheme)</th>
                        <th>Pay And allowance</th>
                        <th>Recurring Contingencies</th>
                        <th>Non-Recurring Contingencies</th>
                        <th>Total Amount</th>
                    </tr>
                    <tr>
                        <td style="text-align: left;">AINP on Agril Acarology<br>(BH.303/2092)</td>
                        <td>{pay_val}</td>
                        <td>{rec_val}</td>
                        <td>{non_rec_val}</td>
                        <td>{tot_val}</td>
                    </tr>
                </table>
                <div>In Rupees: {amt_words}</div>
                <br><br>
                <div>સામેલ: ઉપર મુજબ</div>
                <br><br><br>
                <div style="display: flex; justify-content: space-between;">
                    <div class="bold">પ્રોજેક્ટ ઈન્ચાર્જ</div>
                    <div class="bold right">પ્રાધ્યાપક અને વડા</div>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            st.divider()
            is_approved = st.checkbox("✅ I approve this letter format and content.")
            
            if is_approved:
                # Generate native DOCX Word file based on live edits
                doc_io = generate_comptroller_docx(ref_no, letter_date, body_text, amt_words, pay_amt, total_rec, non_rec_amt, total_amt)
                
                letter_filename = sanitize_filename(
                    f"Letter_to_Comptroller_{selected_inst_data.get('type', 'Installment')}_{selected_inst_data.get('pfms_id', 'UNKNOWN_PFMS')}.docx"
                )
                st.download_button(
                    label="📥 Download Approved Letter (.docx)",
                    data=doc_io,
                    file_name=letter_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.info("No pending PFMS receipts to generate letters for.")

        st.divider()
        st.subheader("Activate Funds (Upload Comptroller Order)")
        st.write("Once received, upload the *'compotrollar grant relased latter'* to activate funds for utilization.")
        
        comp_file = st.file_uploader("Upload Comptroller Office Order PDF", type=['pdf'], key="comp_up")
        
        # 1. Multi-select allows activating multiple installments at once, keyed by unique PFMS ID.
        activation_options = {}
        for inst in data.get('installments', []):
            if "State Share" in inst.get('type', ''):
                continue
            pfms_id = inst.get('pfms_id', '')
            label = f"{inst.get('type', '')} | {pfms_id} | ₹{coerce_amount(inst.get('amount')):,.2f}"
            activation_options[label] = pfms_id
        insts_to_activate = st.multiselect("This order relates to installment(s):", list(activation_options.keys()), key="act_type_multi")
        selected_pfms_ids = {activation_options[label] for label in insts_to_activate}

        # 2. Toggle to automatically calculate and inject the 25% State Share
        add_state_share = st.checkbox("Automatically generate and add the matching 25% State Share", value=True, help="Calculates 25% matching funds for non-TSP heads based on the 75% ICAR PFMS receipt.")

        if comp_file and st.button("Verify and Activate Funds"):
            if not insts_to_activate:
                st.warning("Please select at least one installment to activate.")
            else:
                with st.spinner("Activating Funds and Calculating State Share..."):
                    new_state_shares = []
                    activated_pfms_ids = []
                    safe_order_name = f"{selected_fy}_Comptroller_Order_{datetime.now().strftime('%Y%m%d%H%M%S')}.pdf"
                    save_uploaded_document(comp_file, safe_order_name)
                    existing_pfms_ids = {inst.get('pfms_id') for inst in data.get('installments', [])}
                    
                    for inst in data['installments']:
                        # Check if selected and not already activated to prevent duplication
                        if inst.get('pfms_id') in selected_pfms_ids and not inst.get('comptroller_order_uploaded'):
                            inst['available'] = True
                            inst['comptroller_order_uploaded'] = True
                            inst['comptroller_order_file'] = safe_order_name
                            activated_pfms_ids.append(inst.get('pfms_id'))
                            
                            # Calculate the 25% State Share mathematically
                            if add_state_share:
                                state_heads = {}
                                state_total = 0.0
                                
                                for h, amt in inst.get('heads', {}).items():
                                    # TSP is 100% ICAR, so it gets no state share match
                                    if "TSP" not in h.upper() and amt > 0:
                                        # If ICAR is 75%, then State is 25%. (Therefore State = ICAR / 3)
                                        state_amt = round(float(amt) / 3.0, 2)
                                        state_heads[h] = state_amt
                                        state_total += state_amt
                                
                                # Create a separate ledger entry for the State Share
                                state_pfms_id = "STATE_MATCH_" + inst.get('pfms_id', '')
                                if state_total > 0 and state_pfms_id not in existing_pfms_ids:
                                    state_inst = {
                                        "date": datetime.now().strftime("%Y-%m-%d"),
                                        "quarter": inst.get('quarter', 'Q1'),
                                        "installment_num": inst.get('installment_num', '') + " (State Share)",
                                        "purpose": "State Share 25%",
                                        "pfms_id": state_pfms_id,
                                        "amount": state_total,
                                        "heads": state_heads,
                                        "type": inst.get('type', 'Installment') + " (State Share)",
                                        "available": True,
                                        "comptroller_order_uploaded": True,
                                        "source_pfms_id": inst.get('pfms_id'),
                                        "comptroller_order_file": safe_order_name
                                    }
                                    new_state_shares.append(state_inst)
                                    existing_pfms_ids.add(state_pfms_id)

                    # Append the newly created State Share records safely
                    if new_state_shares:
                        data['installments'].extend(new_state_shares)
                        
                    save_data(
                        data,
                        selected_fy,
                        audit_action="fund_activation",
                        audit_details={
                            "activated_pfms_ids": activated_pfms_ids,
                            "state_share_records": [inst.get('pfms_id') for inst in new_state_shares],
                            "order_file": safe_order_name,
                        },
                    )
                    st.success(f"Funds for {', '.join(activated_pfms_ids)} are now READY! Matching State Share generated.")
                    st.rerun()


# --- TAB 5: MONTHLY SPEND ---
    with tabs[4]:
        st.header("Monthly Expenditure Tracking")
        
        today = date.today()
        month_to_process = st.selectbox("Month", [datetime(2000, m, 1).strftime('%B') for m in range(1, 13)], index=today.month-1)
        year_to_process = st.number_input("Year", value=today.year, min_value=2024, max_value=2030)

        with st.expander("Add New Expenditure Entry", expanded=True):
            with st.form("spend_form", clear_on_submit=True):
                col1, col2 = st.columns(2)
                exp_date = col1.date_input("Expenditure Date")
                exp_head = col1.selectbox("Budget Head", BUDGET_HEADS)
                exp_amt = col2.number_input("Amount Spent (₹)", min_value=0.0)
                exp_detail = col2.text_area("Expenditure Details/Voucher Info")
                
                if st.form_submit_button("Add Entry"):
                    new_exp = {
                        "date": exp_date.strftime("%Y-%m-%d"),
                        "head": exp_head,
                        "detail": exp_detail,
                        "amount": exp_amt
                    }
                    data['expenditure'].append(new_exp)
                    save_data(
                        data,
                        selected_fy,
                        audit_action="expenditure_add",
                        audit_details={"date": new_exp["date"], "head": exp_head, "amount": exp_amt},
                    )
                    st.toast("Expenditure Added and Backed up.")

# Re-initialize df_exp here so both this tab and the chatbot can see it
        exp_list = data.get('expenditure', [])
        
        # Inject a temporary index so we know exactly which item to delete
        for i, exp in enumerate(exp_list):
            exp['_orig_idx'] = i
            
        df_exp = pd.DataFrame(exp_list)
        
        st.subheader(f"Spend list for {month_to_process} {year_to_process}")
        
        if not df_exp.empty:
            df_exp['date_obj'] = pd.to_datetime(df_exp['date'], errors="coerce")
            df_exp = df_exp.dropna(subset=['date_obj'])
            current_month_exp = df_exp[
                (df_exp['date_obj'].dt.strftime('%B') == month_to_process) & 
                (df_exp['date_obj'].dt.year == year_to_process)
            ]
            
            if not current_month_exp.empty:
                # Render Table Header
                hc1, hc2, hc3, hc4, hc5 = st.columns([1.5, 2.5, 1.5, 3, 1])
                hc1.markdown("**Date**")
                hc2.markdown("**Budget Head**")
                hc3.markdown("**Amount (₹)**")
                hc4.markdown("**Details**")
                hc5.markdown("**Action**")
                
                # Render Rows with Delete Buttons
                for _, row in current_month_exp.iterrows():
                    c1, c2, c3, c4, c5 = st.columns([1.5, 2.5, 1.5, 3, 1])
                    c1.write(row['date'])
                    c2.write(row['head'])
                    c3.write(f"₹{float(row['amount']):,.2f}")
                    c4.write(row['detail'])
                    
                    # The unique Delete Button for this specific row
                    if c5.button("🗑️ Delete", key=f"del_exp_{row['_orig_idx']}"):
                        data['expenditure'].pop(row['_orig_idx'])
                        
                        # Clean up the temporary index before saving
                        for e in data['expenditure']: 
                            e.pop('_orig_idx', None) 
                            
                        save_data(
                            data,
                            selected_fy,
                            audit_action="expenditure_delete",
                            audit_details={"date": row['date'], "head": row['head'], "amount": float(row['amount'])},
                        )
                        st.toast("Entry deleted successfully!")
                        st.rerun()
            else:
                st.info(f"No expenditure recorded for {month_to_process} {year_to_process}.")
                
            # Clean up the temporary index from memory so it doesn't affect save files
            for e in data['expenditure']: 
                e.pop('_orig_idx', None)
            
            st.divider()
            
# --- SUMMARY TABLES SECTION ---
            col_sum1, col_sum2 = st.columns(2)
            
            # 1. Monthly Summary Table
            with col_sum1:
                st.markdown(f"**📊 Monthly Summary ({month_to_process} {year_to_process})**")
                
                # Initialize dictionary with all standard heads set to 0
                monthly_summary = {head: 0.0 for head in BUDGET_HEADS}
                
                # If there are expenses this month, update the dictionary
                if not current_month_exp.empty:
                    month_grouped = current_month_exp.groupby('head')['amount'].sum().to_dict()
                    monthly_summary.update(month_grouped)
                
                # Convert to DataFrame for clean display
                df_month_sum = pd.DataFrame(list(monthly_summary.items()), columns=['Budget Head', 'Amount (₹)'])
                
                # Add "Grand Total" Row at the bottom for the Month
                total_month_amount = df_month_sum['Amount (₹)'].sum()
                df_month_sum.loc[len(df_month_sum)] = ['**GRAND TOTAL**', total_month_amount]
                
                st.dataframe(df_month_sum, width="stretch", hide_index=True)

            # 2. Yearly (FY) Summary Table with Grand Total
            with col_sum2:
                st.markdown(f"**📈 Yearly Summary (FY {selected_fy})**")
                
                # Initialize dictionary with all standard heads set to 0
                yearly_summary = {head: 0.0 for head in BUDGET_HEADS}
                
                # Update with actual yearly sums
                year_grouped = df_exp.groupby('head')['amount'].sum().to_dict()
                yearly_summary.update(year_grouped)
                
                df_year_sum = pd.DataFrame(list(yearly_summary.items()), columns=['Budget Head', 'Amount (₹)'])
                
                # Add "Grand Total" Row at the bottom for the Year
                total_amount = df_year_sum['Amount (₹)'].sum()
                df_year_sum.loc[len(df_year_sum)] = ['**GRAND TOTAL**', total_amount]
                
                st.dataframe(df_year_sum, width="stretch", hide_index=True)
            # =====================================================================
            # 👇 PASTE THIS NEW SECTION RIGHT BELOW THE SUMMARY TABLES 👇
            # =====================================================================
            st.divider()
            st.subheader("📜 Cumulative Expenditure Log (Yearly)")
            st.write("View all expenditures, track progressive totals, filter by Budget Head, and delete accidental entries.")
            
            # Filter by Budget Head
            selected_log_head = st.selectbox("Filter by Budget Head:", ["All Heads"] + BUDGET_HEADS, key="log_head_filter")
            
            # Re-inject original index so we can delete safely
            exp_list_yearly = data.get('expenditure', [])
            for i, exp in enumerate(exp_list_yearly):
                exp['_orig_idx'] = i
                
            df_yearly_log = pd.DataFrame(exp_list_yearly)
            
            if not df_yearly_log.empty:
                df_yearly_log['date_obj'] = pd.to_datetime(df_yearly_log['date'], errors="coerce")
                df_yearly_log = df_yearly_log.dropna(subset=['date_obj'])
                
                # Step 1: Filter by selected head
                if selected_log_head != "All Heads":
                    df_yearly_log = df_yearly_log[df_yearly_log['head'] == selected_log_head]
                
                if not df_yearly_log.empty:
                    # Step 2: Sort Oldest to Newest to calculate the Progressive Total correctly
                    df_yearly_log = df_yearly_log.sort_values(by='date_obj', ascending=True)
                    df_yearly_log['progressive_total'] = df_yearly_log['amount'].astype(float).cumsum()
                    
                    # Step 3: Sort Newest to Oldest so the latest entries show up on top
                    df_yearly_log = df_yearly_log.sort_values(by='date_obj', ascending=False)

                    # Show final cumulative total and Download Button side-by-side
                    filtered_total = df_yearly_log['amount'].astype(float).sum()
                    
                    col_tot, col_dl = st.columns([3, 1])
                    with col_tot:
                        st.success(f"**Final Total Spent for {selected_log_head}:** ₹{filtered_total:,.2f}")
                        
                    with col_dl:
                        # Auto-generate PDF in the background
                        pdf = FPDF(orientation='L')
                        pdf.add_page()
                        pdf.set_font("Helvetica", 'B', 14)
                        pdf.cell(
                            0,
                            10,
                            f"Expenditure Log - {selected_log_head} (FY {selected_fy})",
                            new_x=XPos.LMARGIN,
                            new_y=YPos.NEXT,
                            align='C',
                        )
                        pdf.ln(5)
                        
                        pdf.set_font("Helvetica", 'B', 10)
                        cols = [25, 60, 30, 35, 125]
                        headers = ["Date", "Budget Head", "Amount", "Prog. Total", "Details"]
                        for i, h in enumerate(headers):
                            pdf.cell(cols[i], 10, h, border=1, new_x=XPos.RIGHT, new_y=YPos.TOP, align='C')
                        pdf.ln()
                        
                        pdf.set_font("Helvetica", '', 10)
                        for _, r in df_yearly_log.iterrows():
                            pdf.cell(cols[0], 10, str(r['date']), border=1, new_x=XPos.RIGHT, new_y=YPos.TOP, align='C')
                            pdf.cell(cols[1], 10, str(r['head'])[:30], border=1, new_x=XPos.RIGHT, new_y=YPos.TOP, align='L')
                            pdf.cell(cols[2], 10, f"{float(r['amount']):,.2f}", border=1, new_x=XPos.RIGHT, new_y=YPos.TOP, align='R')
                            pdf.cell(cols[3], 10, f"{float(r['progressive_total']):,.2f}", border=1, new_x=XPos.RIGHT, new_y=YPos.TOP, align='R')
                            
                            # Clean details to avoid PDF encoding errors (removes emojis/unsupported chars)
                            detail_txt = str(r['detail']).replace('\n', ' ').encode('latin-1', 'ignore').decode('latin-1')[:75]
                            pdf.cell(cols[4], 10, detail_txt, border=1, new_x=XPos.RIGHT, new_y=YPos.TOP, align='L')
                            pdf.ln()
                            
                        st.download_button(
                            label="📥 Download Log (PDF)",
                            data=pdf_to_bytes(pdf),
                            file_name=sanitize_filename(f"Expenditure_Log_{selected_log_head}_{selected_fy}.pdf".replace(" ", "_")),
                            mime="application/pdf",
                            key="dl_log_pdf"
                        )
                    
                    # Render Table Header (Now with 6 columns)
                    yc1, yc2, yc3, yc4, yc5, yc6 = st.columns([1.2, 2.2, 1.5, 1.5, 2.6, 1])
                    yc1.markdown("**Date**")
                    yc2.markdown("**Budget Head**")
                    yc3.markdown("**Amount (₹)**")
                    yc4.markdown("**Prog. Total (₹)**")
                    yc5.markdown("**Details**")
                    yc6.markdown("**Action**")
                    
                    # Render Rows
                    for _, row in df_yearly_log.iterrows():
                        rc1, rc2, rc3, rc4, rc5, rc6 = st.columns([1.2, 2.2, 1.5, 1.5, 2.6, 1])
                        rc1.write(row['date'])
                        rc2.write(row['head'])
                        rc3.write(f"₹{float(row['amount']):,.2f}")
                        rc4.write(f"**₹{float(row['progressive_total']):,.2f}**")  # <-- Progressive Total Displayed Here
                        rc5.write(row['detail'])
                        
                        # Unique delete button for the yearly log
                        if rc6.button("🗑️ Delete", key=f"del_yr_exp_{row['_orig_idx']}"):
                            data['expenditure'].pop(row['_orig_idx'])
                            
                            # Clean up index before saving
                            for e in data['expenditure']:
                                e.pop('_orig_idx', None)
                                
                            save_data(
                                data,
                                selected_fy,
                                audit_action="expenditure_delete",
                                audit_details={"date": row['date'], "head": row['head'], "amount": float(row['amount'])},
                            )
                            st.toast("Yearly log entry deleted successfully!")
                            st.rerun()
                else:
                    st.info(f"No expenditures found for {selected_log_head} in this financial year.")
            else:
                st.info("No expenditures recorded yet.")
                
            # Final cleanup of the temporary index
            for e in data.get('expenditure', []):
                e.pop('_orig_idx', None)

# --- TAB 6: SOE GENERATION ---
    with tabs[5]:
        st.header("Statement of Expenditure (SOE) Generation")
        
        # 1. Custom function to format numbers as Indian Rupees (e.g., 24,74,691.00)
        def format_inr(number):
            if number == 0: return "-"
            is_neg = number < 0
            num_abs = abs(number)
            s, *d = str(f"{num_abs:.2f}").partition(".")
            r = ",".join([s[x-2:x] for x in range(-3, -len(s), -2)][::-1] + [s[-3:]]) if len(s) > 3 else s
            val = f"{r}{d[0]}{d[1]}"
            return f"(-) {val}" if is_neg else val

        # 2. Controls for Date Selection (SMART LINKED TO FY)
        fy_start_year = int(selected_fy.split('-')[0])
        fy_end_year = fy_start_year + 1
        fy_start_date = datetime(fy_start_year, 4, 1) # April 1st of the selected FY
        
        soe_month = st.selectbox("Select SOE Month", [calendar.month_name[m] for m in range(1, 13)], index=date.today().month-1, key="soe_m")
        month_idx = list(calendar.month_name).index(soe_month)
        
        # SMART YEAR LOGIC: Jan, Feb, Mar belong to the end year of the FY. April-Dec belong to the start year.
        if month_idx in [1, 2, 3]:
            soe_year = fy_end_year
        else:
            soe_year = fy_start_year
            
        st.info(f"📅 **Generating SOE for:** {soe_month} {soe_year} (Automatically linked to FY {selected_fy})")
        
        # Calculate last day of the selected month
        last_day = calendar.monthrange(soe_year, month_idx)[1]
        end_date = datetime(soe_year, month_idx, last_day) # Last day of selected month

# 3. Setup Opening Balances Dictionary
        if 'opening_balances' not in data:
            data['opening_balances'] = {
                "Establishment Charges": 0.0, "TA": 0.0, "Contingencies": 0.0, 
                "TSP": 0.0, "Equipments": 0.0, "Works": 0.0
            }
            
        with st.expander(f"⚙️ Set Opening Balances for FY {selected_fy}"):
            st.write(f"💡 **Tip:** Upload last year's signed AUC in **Tab 6** to auto-fill these balances, or edit them manually below.")
            
            with st.form(f"ob_form_{selected_fy}"):
                cols = st.columns(3)
                new_obs = {}
                for idx, (k, v) in enumerate(data['opening_balances'].items()):
                    # The unique key forces Streamlit to wipe the box clean when the FY changes
                    new_obs[k] = cols[idx % 3].number_input(
                        f"{k} (₹)", 
                        value=float(v), 
                        step=1000.0,
                    )
                if st.form_submit_button("💾 Save Opening Balances"):
                    data['opening_balances'] = new_obs
                    save_data(
                        data,
                        selected_fy,
                        audit_action="opening_balances_update",
                        audit_details={"heads": list(new_obs.keys())},
                    )
                    st.success(f"Opening Balances manually updated for {selected_fy}!")
                    st.rerun()

        st.divider()
        st.markdown(f"<h3 style='text-align: center;'>Statement of Expenditure for the month of {soe_month} {soe_year}</h3>", unsafe_allow_html=True)
        st.markdown("**Name of the Centre:** Navsari")
        st.markdown("**Name of the Scheme:** AICRP/AINP on Agricultural Acarology, NAU, Navsari")

        # 4. Data Gathering & SMART Mapping Logic
        ob = data['opening_balances']
        funds = {k: 0.0 for k in ob.keys()}
        exp = {k: 0.0 for k in ob.keys()}
        
        # Smart Keyword Matcher: Converts any saved name into the strict SOE categories
        def get_smart_soe_head(raw_string):
            if not raw_string: return None
            rs = str(raw_string).upper()
            if "PAY" in rs or "ESTABLISHMENT" in rs: return "Establishment Charges"
            if "TA" in rs or "TRAVELLING" in rs: return "TA"
            if "TSP" in rs: return "TSP"
            if "NON" in rs or "EQUIP" in rs or "WORK" in rs: return "Equipments" 
            if "ORC" in rs or "CONTINGENC" in rs or "RECURRING" in rs: return "Contingencies"
            return None

        # Cumulative Funds Received (From April 1 up to end of selected month)
        for inst in data.get('installments', []):
            inst_date = parse_date(inst.get('date'))
            if not inst_date:
                continue
            if fy_start_date <= inst_date <= end_date:
                for h_name, amt in inst.get('heads', {}).items():
                    soe_head = get_smart_soe_head(h_name)
                    if soe_head in funds:
                        funds[soe_head] += float(amt)

        # Cumulative Expenditure (From April 1 up to end of selected month)
        for e in data.get('expenditure', []):
            e_date = parse_date(e.get('date'))
            if not e_date:
                continue
            if fy_start_date <= e_date <= end_date:
                # Combine head and sub_head just in case, ensuring we catch the keyword
                combined_head = f"{e.get('head', '')} {e.get('sub_head', '')}"
                soe_head = get_smart_soe_head(combined_head)
                    
                if soe_head in exp:
                    exp[soe_head] += float(e.get('amount', 0.0))

        # 5. Build the Data Table Array
        data_table = []
        data_table.append(["", "A. Recurring Contingencies", "", "", "", "", "", ""])

        # Calculate A
        rec_heads = [("1.", "Establishment Charges"), ("2.", "TA"), ("3.", "Contingencies"), ("4.", "TSP")]
        tot_A = [0.0]*5 # [OB, Funds, Exp, ICAR, State]
        for sr, head in rec_heads:
            o = ob[head]; f = funds[head]; e = exp[head]
            i = e * 1.0 if head == "TSP" else e * 0.75
            s = 0.0 if head == "TSP" else e * 0.25
            data_table.append([sr, head, format_inr(o), format_inr(f), format_inr(e), format_inr(i), format_inr(s), format_inr(e)])
            tot_A[0] += o; tot_A[1] += f; tot_A[2] += e; tot_A[3] += i; tot_A[4] += s

        data_table.append(["", "Total - A", format_inr(tot_A[0]), format_inr(tot_A[1]), format_inr(tot_A[2]), format_inr(tot_A[3]), format_inr(tot_A[4]), format_inr(tot_A[2])])
        data_table.append(["", "B. Non Recurring Contingencies", "", "", "", "", "", ""])

        # Calculate B
        non_rec_heads = [("1.", "Equipments"), ("2.", "Works")]
        tot_B = [0.0]*5
        for sr, head in non_rec_heads:
            o = ob[head]; f = funds[head]; e = exp[head]
            i = e * 0.75
            s = e * 0.25
            data_table.append([sr, head, format_inr(o), format_inr(f), format_inr(e), format_inr(i), format_inr(s), format_inr(e)])
            tot_B[0] += o; tot_B[1] += f; tot_B[2] += e; tot_B[3] += i; tot_B[4] += s

        data_table.append(["", "Total - B", format_inr(tot_B[0]), format_inr(tot_B[1]), format_inr(tot_B[2]), format_inr(tot_B[3]), format_inr(tot_B[4]), format_inr(tot_B[2])])

        # Grand Total
        g_tot = [tot_A[j] + tot_B[j] for j in range(5)]
        data_table.append(["", "Grand Total A+B", format_inr(g_tot[0]), format_inr(g_tot[1]), format_inr(g_tot[2]), format_inr(g_tot[3]), format_inr(g_tot[4]), format_inr(g_tot[2])])

        # 6. Render Data
        columns = [
            "Sr. No.", "Head", f"Opening Balance as on 01.04.{fy_start_year}", f"Funds Received from Council", 
            f"Expenditure up to {soe_month} {soe_year}", "75% ICAR Share", "25% State Share", "Total"
        ]
        
        df_soe = pd.DataFrame(data_table, columns=columns)

        st.markdown("💡 **Tip: Your SOE is now calculated automatically. You can still click to edit values if needed before downloading.**")
        edited_df = st.data_editor(df_soe, width="stretch", hide_index=True)
        st.markdown("<ul><li>In 2025-26 State share released only in Pay and allowances</li></ul>", unsafe_allow_html=True)

        st.divider()
        if st.button("Generate SOE Word Document", key="soe_btn_new"):
            with st.spinner("Creating formatted Word file..."):
                soe_doc_buffer = create_word_doc(edited_df, soe_month, soe_year, last_day)
                st.success("SOE Generated with dynamic dates and calculations!")
                
                st.download_button(
                    label="📥 Download SOE Word File",
                    data=soe_doc_buffer,
                    file_name=f"SOE_{soe_month}_{soe_year}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                # =====================================================================
        # 👇 NEW SECTION: 12-MONTH YEARLY SUMMARY TABLE 👇
        # =====================================================================
        st.divider()
        st.markdown(f"<h3 style='text-align: center;'>12-Month Expenditure Summary (FY {selected_fy})</h3>", unsafe_allow_html=True)
        
        # Define the 12 months from April to March
        months_order = ['April', 'May', 'June', 'July', 'August', 'September', 'October', 'November', 'December', 'January', 'February', 'March']
        short_months = [m[:3] for m in months_order] # e.g., 'Apr', 'May' for table columns
        
        # Initialize dictionary to hold 12 months of data for each head
        yearly_exp = {
            "Establishment Charges": {m: 0.0 for m in months_order},
            "TA": {m: 0.0 for m in months_order},
            "Contingencies": {m: 0.0 for m in months_order},
            "TSP": {m: 0.0 for m in months_order},
            "Equipments": {m: 0.0 for m in months_order},
            "Works": {m: 0.0 for m in months_order},
        }
        
        fy_end_date = datetime(fy_end_year, 3, 31) # Ends March 31st

        # Distribute all expenditures from the FY into their exact month buckets
        for e in data.get('expenditure', []):
            e_date = parse_date(e.get('date'))
            if not e_date:
                continue
            if fy_start_date <= e_date <= fy_end_date:
                m_name = calendar.month_name[e_date.month]
                combined_head = f"{e.get('head', '')} {e.get('sub_head', '')}"
                soe_head = get_smart_soe_head(combined_head) # Reuses your smart matcher!
                
                if soe_head in yearly_exp:
                    yearly_exp[soe_head][m_name] += float(e.get('amount', 0.0))

        # Build the 14-column layout for the Yearly Table
        y_table = []
        y_cols = ["Budget Head"] + short_months + ["Total"]

        y_table.append(["A. Recurring Contingencies"] + [""] * 13)

        # Calculate A (Recurring)
        tot_A_y = [0.0] * 13 # 12 months + 1 total
        for head in ["Establishment Charges", "TA", "Contingencies", "TSP"]:
            row = [head]
            row_tot = 0.0
            for i, m in enumerate(months_order):
                val = yearly_exp[head][m]
                row.append(format_inr(val))
                row_tot += val
                tot_A_y[i] += val
            row.append(format_inr(row_tot))
            tot_A_y[12] += row_tot
            y_table.append(row)

        y_table.append(["Total - A"] + [format_inr(v) for v in tot_A_y])
        y_table.append(["B. Non Recurring Contingencies"] + [""] * 13)

        # Calculate B (Non-Recurring)
        tot_B_y = [0.0] * 13
        for head in ["Equipments", "Works"]:
            row = [head]
            row_tot = 0.0
            for i, m in enumerate(months_order):
                val = yearly_exp[head][m]
                row.append(format_inr(val))
                row_tot += val
                tot_B_y[i] += val
            row.append(format_inr(row_tot))
            tot_B_y[12] += row_tot
            y_table.append(row)

        y_table.append(["Total - B"] + [format_inr(v) for v in tot_B_y])

        # Grand Total Calculation
        g_tot_y = [tot_A_y[i] + tot_B_y[i] for i in range(13)]
        y_table.append(["Grand Total A+B"] + [format_inr(v) for v in g_tot_y])

        # Render Yearly Data
        df_yearly = pd.DataFrame(y_table, columns=y_cols)
        
        st.markdown("💡 **Tip: Click to edit values before downloading your 12-Month Summary.**")
        edited_yearly_df = st.data_editor(df_yearly, width="stretch", hide_index=True, key="yearly_summary_editor")

        # Separate Download Button for 12-Month Summary
        if st.button("Generate 12-Month Summary Word Doc", key="yearly_btn_new"):
            with st.spinner("Creating wide landscape Word file..."):
                yearly_doc_buffer = create_yearly_word_doc(edited_yearly_df, selected_fy)
                st.success("12-Month Summary Generated!")
                
                st.download_button(
                    label="📥 Download Yearly Summary Word File",
                    data=yearly_doc_buffer,
                    file_name=f"12_Month_Summary_FY_{selected_fy}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                # =====================================================================
        # 🔵 SECTION 2: ANNUAL SOE (FULL FY)
        # =====================================================================
        st.divider()
        st.markdown(f"<h3 style='text-align: center;'>Annual Statement of Expenditure for FY {selected_fy}</h3>", unsafe_allow_html=True)
        st.info(f"Calculated from 01.04.{fy_start_year} up to 31.03.{fy_end_year}")

        funds_y = {k: 0.0 for k in ob.keys()}
        exp_y = {k: 0.0 for k in ob.keys()}

        # Annual Cumulative Funds
        for inst in data.get('installments', []):
            inst_date = parse_date(inst.get('date'))
            if not inst_date:
                continue
            if fy_start_date <= inst_date <= fy_end_date:
                for h_name, amt in inst.get('heads', {}).items():
                    soe_head = get_smart_soe_head(h_name)
                    if soe_head in funds_y: funds_y[soe_head] += float(amt)

        # Annual Cumulative Exp
        for e in data.get('expenditure', []):
            e_date = parse_date(e.get('date'))
            if not e_date:
                continue
            if fy_start_date <= e_date <= fy_end_date:
                combined_head = f"{e.get('head', '')} {e.get('sub_head', '')}"
                soe_head = get_smart_soe_head(combined_head)
                if soe_head in exp_y: exp_y[soe_head] += float(e.get('amount', 0.0))

        # Build Annual Table Array
        data_table_y = [["", "A. Recurring Contingencies", "", "", "", "", "", ""]]
        tot_A_y = [0.0]*5
        for sr, head in rec_heads:
            o = ob[head]; f = funds_y[head]; e = exp_y[head]
            i = e * 1.0 if head == "TSP" else e * 0.75; s = 0.0 if head == "TSP" else e * 0.25
            data_table_y.append([sr, head, format_inr(o), format_inr(f), format_inr(e), format_inr(i), format_inr(s), format_inr(e)])
            tot_A_y[0] += o; tot_A_y[1] += f; tot_A_y[2] += e; tot_A_y[3] += i; tot_A_y[4] += s

        data_table_y.append(["", "Total - A", format_inr(tot_A_y[0]), format_inr(tot_A_y[1]), format_inr(tot_A_y[2]), format_inr(tot_A_y[3]), format_inr(tot_A_y[4]), format_inr(tot_A_y[2])])
        data_table_y.append(["", "B. Non Recurring Contingencies", "", "", "", "", "", ""])

        tot_B_y = [0.0]*5
        for sr, head in non_rec_heads:
            o = ob[head]; f = funds_y[head]; e = exp_y[head]
            i = e * 0.75; s = e * 0.25
            data_table_y.append([sr, head, format_inr(o), format_inr(f), format_inr(e), format_inr(i), format_inr(s), format_inr(e)])
            tot_B_y[0] += o; tot_B_y[1] += f; tot_B_y[2] += e; tot_B_y[3] += i; tot_B_y[4] += s

        data_table_y.append(["", "Total - B", format_inr(tot_B_y[0]), format_inr(tot_B_y[1]), format_inr(tot_B_y[2]), format_inr(tot_B_y[3]), format_inr(tot_B_y[4]), format_inr(tot_B_y[2])])
        g_tot_y = [tot_A_y[j] + tot_B_y[j] for j in range(5)]
        data_table_y.append(["", "Grand Total A+B", format_inr(g_tot_y[0]), format_inr(g_tot_y[1]), format_inr(g_tot_y[2]), format_inr(g_tot_y[3]), format_inr(g_tot_y[4]), format_inr(g_tot_y[2])])

        columns_y = [
            "Sr. No.", "Head", f"Opening Balance as on 01.04.{fy_start_year}", f"Funds Received during FY {selected_fy}", 
            f"Expenditure during FY {selected_fy}", "75% ICAR Share", "25% State Share", "Total"
        ]
        
        df_soe_y = pd.DataFrame(data_table_y, columns=columns_y)
        edited_df_y = st.data_editor(df_soe_y, width="stretch", hide_index=True, key="annual_edit")

        if st.button("Generate Annual SOE Word Document", key="soe_btn_y"):
            with st.spinner("Creating Annual Word file..."):
                annual_doc_buffer = create_annual_word_doc(edited_df_y, selected_fy)
                st.success("Annual SOE Generated!")
                st.download_button(
                    label="📥 Download Annual SOE Word File",
                    data=annual_doc_buffer,
                    file_name=f"Annual_SOE_FY_{selected_fy}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# =====================================================================
    # 👇 NEW SECTION: TAB 6 - AUC GENERATION 👇
    # =====================================================================
    with tabs[6]:
        st.header("Audit Utilization Certificate (AUC) & Forwarding Letter")
        st.write("Generate the final year-end AUC and its corresponding Gujarati forwarding letter.")
        
# --- NEW: Historical AUC Archive (Dynamic Years) ---
        with st.expander("📂 Historical AUC Archive & Auto-Balance Extraction", expanded=False):
            st.write("Upload previous years' signed AUCs to maintain a digital archive. **If you upload the AUC for the strictly previous year, the AI will automatically extract the closing balances and send them to Tab 5!**")
            
            # Calculate what the "previous year" string should be
            prev_y_start = int(selected_fy.split('-')[0]) - 1
            prev_y_end = str(prev_y_start + 1)[-2:]
            prev_fy_string = f"{prev_y_start}-{prev_y_end}"
            
            # Generate a dynamic list of years (10 years back, 2 years forward)
            current_y = datetime.now().year
            dynamic_years = [f"{y}-{str(y+1)[-2:]}" for y in range(current_y - 10, current_y + 2)]
            
            # Select which year's archive you want to view/upload
            selected_archive_year = st.selectbox(
                "Select Year to Upload or View:", 
                dynamic_years, 
                index=dynamic_years.index(prev_fy_string) if prev_fy_string in dynamic_years else 0
            )
            
            pdf_path = safe_documents_path(f"AUC_Archive_{selected_archive_year}.pdf")
            
            col_up, col_dl = st.columns(2)
            
            with col_up:
                up_file = st.file_uploader(f"Upload AUC for {selected_archive_year}", type=['pdf'], key="up_arc_dyn")
                if up_file and st.button(f"Save & Process {selected_archive_year}", key="btn_ext_dyn"):
                    # 1. Save the file permanently
                    save_uploaded_document(up_file, f"AUC_Archive_{selected_archive_year}.pdf")
                    
                    # 2. If it is the strictly previous year, trigger the AI Extraction!
                    if selected_archive_year == prev_fy_string:
                        with st.spinner(f"Extracting closing balances from {selected_archive_year} to use as opening balances..."):
                            auc_prompt = {
                                "context": "This is a previous year's Audit Utilization Certificate (AUC). Extract the 'Closing balance as on 31st March' for each specific budget head. If a balance is negative (e.g. (-) 1,38,340.60), return it as a negative number (-138340.60).",
                                "structure": {
                                    "Establishment Charges": 0.0, "TA": 0.0, "Contingencies": 0.0,
                                    "TSP": 0.0, "Equipments": 0.0, "Works": 0.0
                                }
                            }
                            extracted_bals = process_upload_with_ai(up_file, auc_prompt)
                            if extracted_bals:
                                if 'opening_balances' not in data:
                                    data['opening_balances'] = {}
                                for k in ["Establishment Charges", "TA", "Contingencies", "TSP", "Equipments", "Works"]:
                                    ext_val = float(extracted_bals.get(k) or 0.0)
                                    data['opening_balances'][k] = ext_val
                                     
                                save_data(
                                    data,
                                    selected_fy,
                                    audit_action="auc_balance_extract",
                                    audit_details={"archive_year": selected_archive_year},
                                )
                                st.success("Balances extracted and sent to Tab 5 SOE!")
                    else:
                        append_audit_log(selected_fy, "auc_archive_upload", {"archive_year": selected_archive_year})
                        st.success(f"Archive saved for {selected_archive_year}!")
                    st.rerun()
                    
            with col_dl:
                st.write("###") # Spacing to align with uploader
                auc_bytes = load_document_bytes(pdf_path)
                if auc_bytes:
                    st.success(f"✅ AUC for {selected_archive_year} is on file.")
                    st.download_button(
                        label=f"📥 Download {selected_archive_year} AUC",
                        data=auc_bytes,
                        file_name=f"AUC_Signed_{selected_archive_year}.pdf",
                        mime="application/pdf",
                        key="dl_arc_dyn",
                    )
                else:
                    st.info("❌ No file uploaded for this year yet.")
        
        st.divider()

        # Calculate Base Variables for the FY (Existing Code continues below...)
        fy_start_year = int(selected_fy.split('-')[0])
        fy_end_year = int("20" + selected_fy.split('-')[1])
        fy_start_date = datetime(fy_start_year, 4, 1)
        fy_end_date = datetime(fy_end_year, 3, 31)
        
        # 1. Gather Installment Data
        inst_data = []
        tot_remittance = 0.0
        for idx, inst in enumerate(data.get('installments', [])):
            inst_date = parse_date(inst.get('date'))
            if not inst_date:
                continue
            if fy_start_date <= inst_date <= fy_end_date:
                amt = float(inst.get('amount', 0.0))
                inst_data.append([idx+1, f"DSC Transaction payment advice report Dated: {inst.get('date')}", f"{amt:,.2f}"])
                tot_remittance += amt
        inst_data.append(["", "Total :-", f"{tot_remittance:,.2f}"])
        
        # 2. Gather Expenditure Data
        def format_inr_auc(number):
            if number == 0: return "-"
            is_neg = number < 0
            num_abs = abs(number)
            s, *d = str(f"{num_abs:.2f}").partition(".")
            r = ",".join([s[x-2:x] for x in range(-3, -len(s), -2)][::-1] + [s[-3:]]) if len(s) > 3 else s
            return f"(-) {r}{d[0]}{d[1]}" if is_neg else f"{r}{d[0]}{d[1]}"

        # Calculate Total ICAR Expense
        ob = data.get('opening_balances', {})
        total_opening_bal_icar = sum(float(v) for v in ob.values())
        
        exp_y = { "Establishment Charges": 0.0, "TA": 0.0, "Contingencies": 0.0, "TSP": 0.0, "Equipments": 0.0, "Works": 0.0 }
        
        def get_smart_soe_head(raw_string):
            rs = str(raw_string).upper()
            if "PAY" in rs or "ESTABLISHMENT" in rs: return "Establishment Charges"
            if "TA" in rs or "TRAVELLING" in rs: return "TA"
            if "TSP" in rs: return "TSP"
            if "NON" in rs or "EQUIP" in rs or "WORK" in rs: return "Equipments" 
            if "ORC" in rs or "CONTINGENC" in rs or "RECURRING" in rs: return "Contingencies"
            return None
            
        for e in data.get('expenditure', []):
            e_date = parse_date(e.get('date'))
            if not e_date:
                continue
            if fy_start_date <= e_date <= fy_end_date:
                combined_head = f"{e.get('head', '')} {e.get('sub_head', '')}"
                soe_head = get_smart_soe_head(combined_head)
                if soe_head in exp_y: exp_y[soe_head] += float(e.get('amount', 0.0))

        # Calculate total ICAR share of expenses
        tot_icar_exp = 0.0
        for head, amt in exp_y.items():
            if head == "TSP": tot_icar_exp += amt * 1.0
            else: tot_icar_exp += amt * 0.75
            
        closing_balance = (total_opening_bal_icar + tot_remittance) - tot_icar_exp
        
        t1_data = [
            format_inr_auc(total_opening_bal_icar),
            format_inr_auc(tot_remittance),
            "0.00",
            format_inr_auc(tot_icar_exp),
            format_inr_auc(closing_balance)
        ]
        
        # Build Table 2
        alloc = data.get('allocation', {})
        def get_alloc(head_kw):
            for k,v in alloc.items():
                if head_kw in k.upper(): return v.get('total', 0.0)
            return 0.0
            
        t2_data = [
            ["Recurring", "", "", "", ""],
            ["Pay and Allowance", format_inr_auc(get_alloc("PAY")), format_inr_auc(exp_y["Establishment Charges"]*0.75), format_inr_auc(exp_y["Establishment Charges"]*0.25), format_inr_auc(exp_y["Establishment Charges"])],
            ["Travelling Allowance", format_inr_auc(get_alloc("TA")), format_inr_auc(exp_y["TA"]*0.75), format_inr_auc(exp_y["TA"]*0.25), format_inr_auc(exp_y["TA"])],
            ["Recurring Contingencies", format_inr_auc(get_alloc("ORC")), format_inr_auc(exp_y["Contingencies"]*0.75), format_inr_auc(exp_y["Contingencies"]*0.25), format_inr_auc(exp_y["Contingencies"])],
            ["HRD", "-", "-", "-", "-"],
            ["Non Recurring Contingencies", "", "", "", ""],
            ["Equipment", format_inr_auc(get_alloc("EQUIP")), format_inr_auc(exp_y["Equipments"]*0.75), format_inr_auc(exp_y["Equipments"]*0.25), format_inr_auc(exp_y["Equipments"])],
            ["Works", "-", "-", "-", "-"],
            ["Vehicle (IT)", "-", "-", "-", "-"],
        ]
        
        tot_alloc = sum(get_alloc(k) for k in ["PAY", "TA", "ORC", "EQUIP"])
        tot_state_exp = sum(exp_y[k]*0.25 for k in exp_y if k != "TSP")
        tot_all_exp = sum(exp_y.values())
        t2_data.append(["Total:-", format_inr_auc(tot_alloc), format_inr_auc(tot_icar_exp), format_inr_auc(tot_state_exp), format_inr_auc(tot_all_exp)])
        
        # --- EDITABLE LIVE PREVIEW SECTION ---
        st.divider()
        st.subheader("👀 AUC Live Preview & Editor")
        st.info("💡 Make any adjustments below before generating the final Word document.")
        
        # 1. Editable Installment Table
        st.markdown("**Received Installments:**")
        df_inst = pd.DataFrame(inst_data, columns=["Sr.No", "Letter No and Date", "Amount"])
        edited_df_inst = st.data_editor(df_inst, width="stretch", hide_index=True, key="auc_edit_inst")
        final_inst_data = edited_df_inst.values.tolist()
        
        # 2. Editable Main Paragraph
        st.markdown("**Form of Utilization Certificate & Audit Utilization Certificate**")
        default_cert_text = f"1. Certified that the out of Rs. {format_inr_auc(tot_remittance)} sanctioned during the year {selected_fy} in favour of Comptroller, NAU, Navsari under this Ministry/Department Letter No. given in the margin and Rs. {format_inr_auc(total_opening_bal_icar)} on account of unspent balance of the previous year, a sum of Rs. {format_inr_auc(tot_icar_exp)} has been Utilized for the purpose of Agril. Acarology Research and remaining unutilized at the end of the year has been surrendered (vide No......Dated......) will be adjusted (to be payable the next year)."
        
        edited_cert_text = st.text_area("Certificate Paragraph 1 (Editable):", value=default_cert_text, height=120)
        st.markdown("2. Certified that I have satisfied myself that the condition on which the expenditure was made have dully fulfilled/are being fulfilled and that I have exercised the following check to see that the money was actually utilized for the purpose for which it was sanctioned.")
        
        # 3. Editable Table 1
        st.markdown("**Table 1: Showing the details of receipt and expenditure figure (in Rupees)**")
        t1_columns = [f"Opening balance as on 1st April {fy_start_year}", f"Remittance Received {selected_fy}", "Receipt", f"ICAR share of Exp during {selected_fy}", f"Closing balance as on 31st March {fy_end_year}"]
        df_t1 = pd.DataFrame([t1_data], columns=t1_columns)
        edited_df_t1 = st.data_editor(df_t1, width="stretch", hide_index=True, key="auc_edit_t1")
        final_t1_data = edited_df_t1.values.tolist()[0]
        
        # 4. Editable Table 2
        st.markdown("**Table 2: Showing the head wise details of expenditure figure (in Rupees)**")
        t2_columns = ["Head", f"Allocation for the Year {selected_fy} (100%)", "ICAR share of Expenditure (75%)", "State Share (25%)", "Total Expenditure"]
        df_t2 = pd.DataFrame(t2_data, columns=t2_columns)
        edited_df_t2 = st.data_editor(df_t2, width="stretch", hide_index=True, key="auc_edit_t2")
        final_t2_data = edited_df_t2.values.tolist()
        
        st.divider()

        # --- DOWNLOAD BUTTONS SECTION ---
        col_a, col_b = st.columns(2)
        with col_a:
            st.subheader("1. Audit Utilization Certificate (AUC)")
            st.write("Generates using the edited data from the preview above.")
            if st.button("📥 Generate & Download AUC Document"):
                with st.spinner("Generating AUC..."):
                    # Pass the final edited values into the generator
                    auc_doc = generate_auc_certificate(final_inst_data, final_t1_data, final_t2_data, edited_cert_text, selected_fy)
                    st.download_button("Download AUC (.docx)", data=auc_doc, file_name=f"AUC_FY_{selected_fy}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    
        with col_b:
            st.subheader("2. AUC Forwarding Letter")
            
            # Using the exact defaults from the uploaded image
            ref_no = st.text_input("Reference No. (જા.નં. એસીએન/એન્ટો/એઆઈએનપી-એએ/___/૨૦૨૪):", value="288", key="auc_ref_no")
            letter_date = st.text_input("Date (તારીખ):", value=datetime.now().strftime("%d/%m/%Y"), key="auc_letter_date")
            
            subj = st.text_area("Subject (વિષય):", value=f"AINP on Agricultural Acarology (B.H. 303/2092) નું વર્ષ {selected_fy} નુ Audit Utilization Certificate (AUC) મોકલવા બાબત.", key="auc_subject")
            
            default_body = f"જય ભારત સહ ઉપરોક્ત વિષય અન્વયે સવિનય જણાવવાનું કે, અત્રેના કીટકશાસ્ત્ર વિભાગ ખાતે ચાલતી આઈ.સી.એ.આર. યોજના AINP on Agricultural Acarology (B.H. 303/2092) નું વર્ષ {selected_fy} નુ Audit Utilization Certificate (AUC) આ સાથે સામેલ રાખી મોકલી આપીએ છીએ.\n\nજે આપ સાહેબને વિદિત થાય."
            body = st.text_area("Body Text:", value=default_body, height=150, key="auc_body_text")
            
            if st.button("📥 Generate & Download Forwarding Letter", key="auc_download_btn"):
                with st.spinner("Generating Letter..."):
                    fw_doc = generate_auc_forwarding_docx(ref_no, letter_date, subj, body)
                    st.download_button("Download Forwarding Letter (.docx)", data=fw_doc, file_name=f"AUC_Forwarding_Letter_{selected_fy}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="auc_final_dl_btn")

# --- TAB 7: AI CHATBOT ---
    with tabs[7]:
        st.header("🧠 Advanced Grant Assistant & Report Generator")
        st.write("Ask questions about your grant, or **upload University guidelines, Excel data, or PDFs** to extract information and automatically generate reports.")

        # 1. Temporary Context Uploader
        with st.expander("📂 Upload Documents for Context (PDF, Excel, CSV, Word)", expanded=False):
            st.info("Upload files here, then ask the chatbot to analyze them, summarize rules, or generate a specific report based on their contents.")
            chat_files = st.file_uploader("Upload reference files", type=['pdf', 'xlsx', 'csv', 'docx'], accept_multiple_files=True, key="chat_uploader")

        # 2. Persistent Knowledge Store
        knowledge_metadata = load_ai_knowledge()
        with st.expander("Grant Knowledge Store", expanded=False):
            if genai_client is None:
                st.info("Add GEMINI_API_KEY in Streamlit secrets to create or query the Google File Search knowledge store.")
            else:
                store_name = knowledge_metadata.get("store_name")
                st.caption(f"Store: {store_name or 'Not created yet'}")
                knowledge_files = st.file_uploader(
                    "Index selected grant documents",
                    type=['pdf', 'xlsx', 'csv', 'docx'],
                    accept_multiple_files=True,
                    key="knowledge_store_uploader",
                )
                if knowledge_files and st.button("Upload to Knowledge Store", key="upload_knowledge_store"):
                    uploaded_names = []
                    with st.spinner("Uploading and indexing selected files..."):
                        for knowledge_file in knowledge_files:
                            upload_file_to_knowledge_store(knowledge_file)
                            uploaded_names.append(knowledge_file.name)
                    append_audit_log(selected_fy, "ai_knowledge_upload", {"files": uploaded_names})
                    st.success("Knowledge Store updated.")
                    st.rerun()

                documents = knowledge_metadata.get("documents", [])
                if documents:
                    st.dataframe(pd.DataFrame(documents)[["file_name", "mime_type", "uploaded_at"]], width="stretch", hide_index=True)
                else:
                    st.info("No documents indexed yet.")

        # 3. Controlled Learning Memory
        with st.expander("Controlled Learning Memory", expanded=False):
            memories = load_learning_memories()
            with st.form("add_learning_memory_form"):
                mem_col1, mem_col2 = st.columns(2)
                mem_title = mem_col1.text_input("Title")
                mem_keywords = mem_col2.text_input("Keywords, comma separated")
                mem_text = st.text_area("Memory text")
                if st.form_submit_button("Save Learning Memory"):
                    if not mem_title.strip() or not mem_text.strip():
                        st.error("Please enter a title and memory text.")
                    else:
                        memories.append({
                            "id": datetime.now().strftime("%Y%m%d%H%M%S%f"),
                            "title": mem_title.strip(),
                            "keywords": mem_keywords.strip(),
                            "memory_text": mem_text.strip(),
                            "enabled": True,
                            "created_at": datetime.now().isoformat(timespec="seconds"),
                        })
                        save_learning_memories(memories)
                        append_audit_log(selected_fy, "learning_memory_add", {"title": mem_title.strip()})
                        st.success("Learning memory saved.")
                        st.rerun()

            if memories:
                st.write("Saved memories")
                for memory in memories:
                    status = "Enabled" if memory.get("enabled", True) else "Disabled"
                    with st.container(border=True):
                        st.markdown(f"**{memory.get('title', '')}** ({status})")
                        st.caption(memory.get("keywords", ""))
                        st.write(memory.get("memory_text", ""))
                        col_mem1, col_mem2 = st.columns(2)
                        if col_mem1.button("Toggle Enabled", key=f"toggle_memory_{memory.get('id')}"):
                            memory["enabled"] = not memory.get("enabled", True)
                            save_learning_memories(memories)
                            st.rerun()
                        if col_mem2.button("Delete", key=f"delete_memory_{memory.get('id')}"):
                            memories = [m for m in memories if m.get("id") != memory.get("id")]
                            save_learning_memories(memories)
                            append_audit_log(selected_fy, "learning_memory_delete", {"title": memory.get("title", "")})
                            st.rerun()
            else:
                st.info("No controlled memories saved yet.")

        # 4. Financial Context Setup
        budget_summary = data['revised_allocation'] if data['revised_allocation'] else data['allocation']
        received_summary = sum(coerce_amount(inst.get('amount')) for inst in data['installments'])
        
        df_exp = pd.DataFrame(data.get('expenditure', []))
        spend_summary = {}
        if not df_exp.empty:
            spend_summary = df_exp.groupby('head')['amount'].sum().to_dict()
            
        system_context = f"""
        You are an elite financial assistant and report writer for Dr. Vaibhav Chaudhari, managing the AINP on Agricultural Acarology grant at NAU Navsari.
        The current Financial Year is {selected_fy}.
        Total Funds Received from ICAR: ₹{received_summary:,}
        Budget Allocation Summary: {json.dumps(budget_summary)}
        Current Expenditure by Head: {json.dumps(spend_summary)}
        
        INSTRUCTIONS:
        1. Answer user questions robustly, calculating remaining balances where needed (Allocation - Spend).
        2. If the user asks to generate a report or summary, format it highly professionally using Markdown tables, bold headers, and bullet points.
        3. If the user has provided uploaded files, prioritize extracting the exact rules, guidelines, or data they request from those files.
        4. If LEARNED MEMORY is provided, use it as explicit saved preference/correction.
        5. If File Search sources are available, ground answers in those documents and mention when the answer is based on indexed knowledge.
        """

        # 5. Initialize Chat History in Session State
        if "messages" not in st.session_state:
            st.session_state.messages = []

        # Display UI Chat History
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

        use_file_search = bool(knowledge_metadata.get("store_name")) and st.checkbox(
            "Use Grant Knowledge Store for answers",
            value=bool(knowledge_metadata.get("store_name")),
            key="use_file_search_for_chat",
        )

        # 6. Handle User Prompt
        if prompt := st.chat_input("Ask about grant status, or type 'Generate a report based on the uploaded file' ..."):
            
            # Display user message
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)

            with st.chat_message("assistant"):
                message_placeholder = st.empty()
                
                with st.spinner("Analyzing data and generating response..."):
                    try:
                        if genai_client is None:
                            raise RuntimeError("Please add GEMINI_API_KEY in Streamlit secrets to use the AI assistant.")

                        recent_history = "\n".join(
                            f"{msg['role'].upper()}: {msg['content']}"
                            for msg in st.session_state.messages[-10:-1]
                        )
                        learned = matching_learning_memories(prompt)
                        learned_context = "\n".join(
                            f"- {m.get('title')}: {m.get('memory_text')}"
                            for m in learned
                        )

                        # --- Prepare the current message payload ---
                        user_payload = f"""
                        RECENT CHAT HISTORY:
                        {recent_history or 'No prior chat history.'}

                        LEARNED MEMORY:
                        {learned_context or 'No matching saved learning memory.'}

                        USER PROMPT:
                        {prompt}
                        """
                        message_parts = [types.Part.from_text(text=user_payload)]
                        
                        # --- SMART FILE PROCESSING ---
                        if chat_files:
                            for file in chat_files:
                                ext = file.name.split('.')[-1].lower()
                                
                                # Process PDFs natively through Gemini
                                if ext == 'pdf':
                                    message_parts.append(types.Part.from_bytes(data=file.getvalue(), mime_type="application/pdf"))
                                    
                                # Extract text from Word Docs using python-docx
                                elif ext == 'docx':
                                    try:
                                        doc = Document(io.BytesIO(file.getvalue()))
                                        doc_text = "\n".join([p.text for p in doc.paragraphs])
                                        message_parts.append(types.Part.from_text(text=f"\n--- Content of {file.name} ---\n{doc_text}\n---"))
                                    except Exception:
                                        message_parts.append(types.Part.from_text(text=f"[Could not extract text from Word Doc: {file.name}]"))
                                        
                                # Extract data from Excel/CSV using Pandas
                                elif ext in ['xlsx', 'csv']:
                                    try:
                                        df = pd.read_csv(file) if ext == 'csv' else pd.read_excel(file)
                                        csv_string = df.to_csv(index=False)
                                        message_parts.append(types.Part.from_text(text=f"\n--- Tabular Data from {file.name} ---\n{csv_string}\n---"))
                                    except Exception:
                                        message_parts.append(types.Part.from_text(text=f"[Could not extract table data from {file.name}]"))

                        tools = []
                        if use_file_search:
                            tools.append(types.Tool(file_search=types.FileSearch(
                                file_search_store_names=[knowledge_metadata["store_name"]],
                                top_k=5,
                            )))

                        response = genai_client.models.generate_content(
                            model=GEMINI_CHAT_MODEL,
                            contents=message_parts,
                            config=types.GenerateContentConfig(
                                system_instruction=system_context,
                                tools=tools or None,
                            ),
                        )
                        full_response = response.text
                        citations = extract_file_search_citations(response) if use_file_search else []
                        
                    except Exception as e:
                        full_response = f"Sorry, the AI service encountered an error. Error details: {e}"
                        citations = []

                # Stream the response to the UI
                message_placeholder.markdown(full_response)
                if citations:
                    with st.expander("File Search Sources", expanded=False):
                        st.json(citations[:8])
            
            # Save assistant response to memory
            st.session_state.messages.append({"role": "assistant", "content": full_response})
# --- 4. RUN APPLICATION ---
if __name__ == "__main__":
    if NAU_LOGO and ICAR_LOGO:
        main()
    else:
        st.error("Missing Logos. Please ensure 'nau_logo.png' and 'icar_logo.png' are in the 'logos/' folder.")
        st.write(f"Paths checked: {NAU_LOGO}, {ICAR_LOGO}")
